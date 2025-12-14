"""
Job Matcher Backend - COMPLETE VERSION
With improved error handling and simplified RapidAPI queries
Integrated with CareerLens features for enhanced functionality
"""

import os
import re
import time
import json
import hashlib
from io import BytesIO
from typing import Dict, List, Optional, Tuple
from datetime import datetime
import requests
import streamlit as st
import sqlite3

# Lazy imports for heavy ML/embedding libraries - loaded only when needed
_SentenceTransformer = None
_Pinecone = None
_ServerlessSpec = None


def _get_sentence_transformer_class():
    """Lazy load SentenceTransformer class"""
    global _SentenceTransformer
    if _SentenceTransformer is None:
        from sentence_transformers import SentenceTransformer
        _SentenceTransformer = SentenceTransformer
    return _SentenceTransformer


def _get_pinecone_classes():
    """Lazy load Pinecone classes"""
    global _Pinecone, _ServerlessSpec
    if _Pinecone is None:
        from pinecone import Pinecone, ServerlessSpec
        _Pinecone = Pinecone
        _ServerlessSpec = ServerlessSpec
    return _Pinecone, _ServerlessSpec


# Lazy imports for document processing - loaded when needed
_docx = None
_Document = None
_PyPDF2 = None


def _get_docx():
    """Lazy load docx and related"""
    global _docx, _Document
    if _docx is None:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        _docx = docx
        _Document = Document
    return _docx


def _get_pypdf2():
    """Lazy load PyPDF2"""
    global _PyPDF2
    if _PyPDF2 is None:
        import PyPDF2
        _PyPDF2 = PyPDF2
    return _PyPDF2


# Import docx components needed at function level
# These will be imported when first used
import openai
from openai import AzureOpenAI
from config import Config

# Lazy imports for data processing
_pd = None
_np = None


def _get_pandas():
    """Lazy load pandas"""
    global _pd
    if _pd is None:
        import pandas as pd
        _pd = pd
    return _pd


def _get_numpy():
    """Lazy load numpy"""
    global _np
    if _np is None:
        import numpy as np
        _np = np
    return _np

# Optional imports for enhanced features
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    TIKTOKEN_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor, black
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

# Initialize config
Config.setup()

# ============================================================================
# CACHED MODEL LOADING - Critical for performance
# ============================================================================

# Cache the SentenceTransformer model to avoid reloading on every Streamlit rerun
_cached_model = None

def get_sentence_transformer_model():
    """Get cached SentenceTransformer model - only loads once"""
    global _cached_model
    if _cached_model is None:
        print("📦 Loading sentence transformer model (first time only)...")
        _cached_model = SentenceTransformer(Config.MODEL_NAME)
        print("✅ Model loaded and cached!")
    return _cached_model

# Cache Pinecone client
_cached_pinecone = None

def get_pinecone_client():
    """Get cached Pinecone client - only initializes once"""
    global _cached_pinecone
    if _cached_pinecone is None:
        print("📦 Initializing Pinecone client (first time only)...")
        _cached_pinecone = Pinecone(api_key=Config.PINECONE_API_KEY)
        print("✅ Pinecone client cached!")
    return _cached_pinecone

# Cache LinkedInJobSearcher
_cached_linkedin_searcher = None

def get_linkedin_job_searcher():
    """Get cached LinkedInJobSearcher - only initializes once"""
    global _cached_linkedin_searcher
    if _cached_linkedin_searcher is None:
        print("📦 Initializing LinkedIn Job Searcher (first time only)...")
        # Import here to avoid circular dependency, searcher defined later in file
        _cached_linkedin_searcher = LinkedInJobSearcher(Config.RAPIDAPI_KEY)
        print("✅ LinkedIn Job Searcher cached!")
    return _cached_linkedin_searcher


# ============================================================================
# CAREERLENS UTILITY CLASSES AND FUNCTIONS
# ============================================================================

class TokenUsageTracker:
    """Tracks token usage and costs for API calls (from CareerLens)"""
    def __init__(self):
        self.total_tokens = 0
        self.total_prompt_tokens = 0
        self.total_completion_tokens = 0
        self.total_embedding_tokens = 0
        self.cost_usd = 0.0
        self.embedding_cost_per_1k = 0.00002
        self.gpt4_mini_prompt_cost_per_1k = 0.00015
        self.gpt4_mini_completion_cost_per_1k = 0.0006
    
    def add_embedding_tokens(self, tokens):
        """Track embedding token usage"""
        self.total_embedding_tokens += tokens
        self.total_tokens += tokens
        self.cost_usd += (tokens / 1000) * self.embedding_cost_per_1k
    
    def add_completion_tokens(self, prompt_tokens, completion_tokens):
        """Track completion token usage"""
        self.total_prompt_tokens += prompt_tokens
        self.total_completion_tokens += completion_tokens
        self.total_tokens += prompt_tokens + completion_tokens
        self.cost_usd += (prompt_tokens / 1000) * self.gpt4_mini_prompt_cost_per_1k
        self.cost_usd += (completion_tokens / 1000) * self.gpt4_mini_completion_cost_per_1k
    
    def get_summary(self):
        """Get usage summary"""
        return {
            'total_tokens': self.total_tokens,
            'embedding_tokens': self.total_embedding_tokens,
            'prompt_tokens': self.total_prompt_tokens,
            'completion_tokens': self.total_completion_tokens,
            'estimated_cost_usd': round(self.cost_usd, 4)
        }
    
    def reset(self):
        """Reset counters"""
        self.total_tokens = 0
        self.total_prompt_tokens = 0
        self.total_completion_tokens = 0
        self.total_embedding_tokens = 0
        self.cost_usd = 0.0


class RateLimiter:
    """Simple rate limiter for API calls (from CareerLens)"""
    def __init__(self, max_requests_per_minute=10):
        self.max_requests_per_minute = max_requests_per_minute
        self.request_times = []
    
    def wait_if_needed(self):
        """Wait if rate limit exceeded"""
        if self.max_requests_per_minute <= 0:
            return
        
        now = time.time()
        one_minute_ago = now - 60
        self.request_times = [t for t in self.request_times if t > one_minute_ago]
        
        if len(self.request_times) >= self.max_requests_per_minute:
            oldest_request = min(self.request_times)
            wait_time = 60 - (now - oldest_request) + 1
            if wait_time > 0:
                print(f"⏳ Rate limiting: waiting {wait_time:.1f}s...")
                time.sleep(wait_time)
                now = time.time()
                one_minute_ago = now - 60
                self.request_times = [t for t in self.request_times if t > one_minute_ago]
        
        self.request_times.append(time.time())


def api_call_with_retry(request_func, max_retries=3, initial_delay=1):
    """Execute API call with exponential backoff retry (from CareerLens)"""
    delay = initial_delay
    last_response = None
    
    for attempt in range(max_retries):
        try:
            response = request_func()
            last_response = response
            
            if response.status_code == 429:
                retry_after = response.headers.get('Retry-After', delay * 2)
                try:
                    wait_time = int(retry_after)
                except (ValueError, TypeError):
                    wait_time = delay * 2
                
                print(f"⚠️ Rate limited (attempt {attempt + 1}/{max_retries}), waiting {wait_time}s...")
                time.sleep(wait_time)
                delay = min(delay * 2, 60)
                continue
            
            return response
            
        except requests.exceptions.RequestException as e:
            print(f"⚠️ Request error (attempt {attempt + 1}/{max_retries}): {e}")
            if attempt < max_retries - 1:
                time.sleep(delay)
                delay = min(delay * 2, 60)
            continue
    
    return last_response


def extract_salary_from_text(text):
    """Extract salary information from job description using LLM (from CareerLens)"""
    if not text:
        return None, None
    
    text_for_extraction = text[:3000] if len(text) > 3000 else text
    
    try:
        # Check if API keys are configured - if not, fall back to regex
        is_configured, _ = Config.check_azure_credentials()
        if not is_configured:
            return extract_salary_from_text_regex(text)
        
        client = AzureOpenAI(
            azure_endpoint=Config.AZURE_ENDPOINT,
            api_key=Config.AZURE_API_KEY,
            api_version=Config.AZURE_API_VERSION
        )
        
        prompt = f"""Extract salary information from this job description text. 
Look for salary ranges, amounts, and compensation details. Normalize everything to monthly HKD (Hong Kong Dollars).

JOB DESCRIPTION TEXT:
{text_for_extraction}

Extract and return salary information as JSON with this structure:
{{
    "min_salary_hkd_monthly": <number or null>,
    "max_salary_hkd_monthly": <number or null>,
    "found": true/false,
    "raw_text": "the exact salary text found in the description"
}}

Rules:
- Convert all amounts to monthly HKD (multiply annual by 12, weekly by 4.33, daily by 22)
- If only one amount is found, set both min and max to that value
- If no salary is found, set "found": false and return null for min/max
- Always return valid JSON, no extra explanation"""

        response = client.chat.completions.create(
            model=Config.AZURE_MODEL,
            messages=[
                {"role": "system", "content": "You are a salary extraction expert. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=300,
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        content = response.choices[0].message.content
        salary_data = json.loads(content)
        
        if salary_data.get('found', False):
            min_sal = salary_data.get('min_salary_hkd_monthly')
            max_sal = salary_data.get('max_salary_hkd_monthly')
            if min_sal is not None and max_sal is not None:
                return int(min_sal), int(max_sal)
            elif min_sal is not None:
                return int(min_sal), int(min_sal * 1.2)
        
        return extract_salary_from_text_regex(text)
        
    except Exception as e:
        return extract_salary_from_text_regex(text)


def extract_salary_from_text_regex(text):
    """Fallback regex-based salary extraction (from CareerLens)"""
    if not text:
        return None, None
    
    patterns = [
        r'HKD\s*\$?\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*[-–—]\s*\$?\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)',
        r'(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*[-–—]\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*HKD',
        r'HKD\s*\$?\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*(?:per month|/month|/mth|monthly)',
        r'(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*HKD\s*(?:per month|/month|/mth|monthly)',
        r'\$\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)\s*[-–—]\s*\$?\s*(\d{1,3}(?:,\d{3})*(?:k|K)?)',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            match = matches[0]
            if isinstance(match, tuple) and len(match) == 2:
                min_sal = match[0].replace(',', '').replace('k', '000').replace('K', '000')
                max_sal = match[1].replace(',', '').replace('k', '000').replace('K', '000')
                try:
                    return int(min_sal), int(max_sal)
                except (ValueError, TypeError):
                    pass
    
    return None, None


def calculate_salary_band(matched_jobs):
    """Calculate estimated salary band from matched jobs (from CareerLens)"""
    salaries = []
    
    for result in matched_jobs:
        job = result.get('job', result)
        salary_str = job.get('salary', '')
        if salary_str and salary_str != 'Not specified':
            min_sal, max_sal = extract_salary_from_text(salary_str)
            if min_sal and max_sal:
                salaries.append((min_sal, max_sal))
        
        description = job.get('description', '')
        if description:
            min_sal, max_sal = extract_salary_from_text_regex(description[:5000])
            if min_sal and max_sal:
                salaries.append((min_sal, max_sal))
    
    if not salaries:
        return 45000, 55000
    
    # Use Python's built-in for simple mean calculation instead of numpy
    avg_min = int(sum(s[0] for s in salaries) / len(salaries))
    avg_max = int(sum(s[1] for s in salaries) / len(salaries))
    
    return avg_min, avg_max


def filter_jobs_by_domains(jobs, target_domains):
    """Filter jobs by target domains/industries (from CareerLens)"""
    if not target_domains:
        return jobs
    
    domain_keywords = {
        'FinTech': ['fintech', 'financial technology', 'blockchain', 'crypto', 'payment', 'banking technology', 'digital banking'],
        'ESG & Sustainability': ['esg', 'sustainability', 'environmental', 'green', 'carbon', 'climate', 'renewable'],
        'Data Analytics': ['data analytics', 'data analysis', 'business intelligence', 'bi', 'data science', 'analytics', 'big data'],
        'Digital Transformation': ['digital transformation', 'digitalization', 'digital strategy', 'innovation'],
        'Investment Banking': ['investment banking', 'ib', 'm&a', 'mergers', 'acquisitions', 'capital markets', 'equity research'],
        'Consulting': ['consulting', 'consultant', 'advisory', 'strategy consulting', 'management consulting'],
        'Technology': ['software', 'technology', 'tech', 'engineering', 'developer', 'programming', 'it'],
        'Healthcare': ['healthcare', 'medical', 'health', 'hospital', 'clinical', 'pharmaceutical', 'biotech'],
        'Education': ['education', 'teaching', 'academic', 'university', 'school', 'e-learning', 'edtech'],
        'Real Estate': ['real estate', 'property', 'realty', 'property management'],
        'Retail & E-commerce': ['retail', 'e-commerce', 'ecommerce', 'online retail'],
        'Marketing & Advertising': ['marketing', 'advertising', 'brand', 'digital marketing', 'social media'],
        'Legal': ['legal', 'law', 'attorney', 'lawyer', 'compliance', 'regulatory'],
        'Human Resources': ['human resources', 'hr', 'recruitment', 'talent acquisition', 'people operations'],
        'Operations': ['operations', 'supply chain', 'logistics', 'procurement']
    }
    
    filtered = []
    for job in jobs:
        title_lower = job.get('title', '').lower()
        desc_lower = job.get('description', '').lower()
        company_lower = job.get('company', '').lower()
        combined = f"{title_lower} {desc_lower} {company_lower}"
        
        for domain in target_domains:
            keywords = domain_keywords.get(domain, [domain.lower()])
            if any(keyword.lower() in combined for keyword in keywords):
                filtered.append(job)
                break
    
    return filtered if filtered else jobs


def filter_jobs_by_salary(jobs, min_salary):
    """Filter jobs by minimum salary expectation (from CareerLens)"""
    if not min_salary or min_salary <= 0:
        return jobs
    
    filtered = []
    jobs_without_salary = []
    
    for job in jobs:
        salary_str = job.get('salary', '')
        description = job.get('description', '')
        
        min_sal, max_sal = extract_salary_from_text_regex(salary_str)
        
        if not min_sal:
            min_sal, max_sal = extract_salary_from_text_regex(description)
        
        if min_sal:
            if min_sal >= min_salary or (max_sal and max_sal >= min_salary):
                filtered.append(job)
        else:
            jobs_without_salary.append(job)
    
    if filtered:
        return filtered
    elif jobs_without_salary:
        return jobs_without_salary
    else:
        return []


# ============================================================================
# RESUME FORMATTERS (from CareerLens)
# ============================================================================

def set_cell_shading(cell, color):
    """Set background color for a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_horizontal_line(doc, color="2B5797"):
    """Add a horizontal line to the document"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx_from_json(resume_data, filename="resume.docx"):
    """Generate a modern professional .docx file from structured resume JSON (from CareerLens)"""
    try:
        # Lazy load docx components when needed
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)
        
        PRIMARY_COLOR = RGBColor(43, 87, 151)
        SECONDARY_COLOR = RGBColor(80, 80, 80)
        ACCENT_COLOR = RGBColor(0, 120, 212)
        
        header = resume_data.get('header', {})
        
        # Name Header
        if header.get('name'):
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(header['name'].upper())
            name_run.font.size = Pt(24)
            name_run.font.bold = True
            name_run.font.color.rgb = PRIMARY_COLOR
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_para.paragraph_format.space_after = Pt(4)
        
        # Professional Title
        if header.get('title'):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(header['title'])
            title_run.font.size = Pt(13)
            title_run.font.color.rgb = SECONDARY_COLOR
            title_run.font.italic = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(8)
        
        # Contact Info
        contact_items = []
        if header.get('email'):
            contact_items.append(f"✉ {header['email']}")
        if header.get('phone'):
            contact_items.append(f"📞 {header['phone']}")
        if header.get('location'):
            contact_items.append(f"📍 {header['location']}")
        if header.get('linkedin'):
            contact_items.append(f"💼 {header['linkedin']}")
        
        if contact_items:
            contact_para = doc.add_paragraph()
            contact_text = '  •  '.join(contact_items)
            contact_run = contact_para.add_run(contact_text)
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = SECONDARY_COLOR
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(12)
        
        add_horizontal_line(doc, "2B5797")
        
        # Professional Summary
        if resume_data.get('summary'):
            summary_header = doc.add_paragraph()
            header_run = summary_header.add_run('PROFESSIONAL SUMMARY')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(resume_data['summary'])
            summary_run.font.size = Pt(10)
            summary_run.font.color.rgb = SECONDARY_COLOR
        
        # Key Skills
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            skills_header = doc.add_paragraph()
            header_run = skills_header.add_run('KEY SKILLS')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            skills_para = doc.add_paragraph()
            for i, skill in enumerate(skills):
                skill_run = skills_para.add_run(f" {skill} ")
                skill_run.font.size = Pt(9)
                skill_run.font.color.rgb = PRIMARY_COLOR
                if i < len(skills) - 1:
                    separator = skills_para.add_run("  |  ")
                    separator.font.size = Pt(9)
                    separator.font.color.rgb = RGBColor(180, 180, 180)
        
        # Professional Experience
        experience = resume_data.get('experience', [])
        if experience:
            exp_header = doc.add_paragraph()
            header_run = exp_header.add_run('PROFESSIONAL EXPERIENCE')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            for exp in experience:
                job_header = doc.add_paragraph()
                
                if exp.get('title'):
                    title_run = job_header.add_run(exp['title'])
                    title_run.font.size = Pt(11)
                    title_run.font.bold = True
                    title_run.font.color.rgb = RGBColor(50, 50, 50)
                
                if exp.get('company'):
                    company_run = job_header.add_run(f"  |  {exp['company']}")
                    company_run.font.size = Pt(10)
                    company_run.font.color.rgb = ACCENT_COLOR
                
                if exp.get('dates'):
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(exp['dates'])
                    date_run.font.size = Pt(9)
                    date_run.font.italic = True
                    date_run.font.color.rgb = SECONDARY_COLOR
                
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet and bullet.strip():
                        bullet_para = doc.add_paragraph()
                        bullet_run = bullet_para.add_run("▸  ")
                        bullet_run.font.size = Pt(9)
                        bullet_run.font.color.rgb = ACCENT_COLOR
                        
                        text_run = bullet_para.add_run(bullet.strip())
                        text_run.font.size = Pt(10)
                        text_run.font.color.rgb = SECONDARY_COLOR
                        bullet_para.paragraph_format.left_indent = Inches(0.25)
        
        # Education
        if resume_data.get('education'):
            edu_header = doc.add_paragraph()
            header_run = edu_header.add_run('EDUCATION')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            edu_para = doc.add_paragraph()
            edu_run = edu_para.add_run(resume_data['education'])
            edu_run.font.size = Pt(10)
            edu_run.font.color.rgb = SECONDARY_COLOR
        
        # Certifications
        if resume_data.get('certifications'):
            cert_header = doc.add_paragraph()
            header_run = cert_header.add_run('CERTIFICATIONS & ACHIEVEMENTS')
            header_run.font.size = Pt(11)
            header_run.font.bold = True
            header_run.font.color.rgb = PRIMARY_COLOR
            
            cert_para = doc.add_paragraph()
            cert_run = cert_para.add_run(resume_data['certifications'])
            cert_run.font.size = Pt(10)
            cert_run.font.color.rgb = SECONDARY_COLOR
        
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
        
    except Exception as e:
        st.error(f"Error generating DOCX: {e}")
        return None


def generate_pdf_from_json(resume_data, filename="resume.pdf"):
    """Generate a professional PDF file from structured resume JSON (from CareerLens)"""
    if not REPORTLAB_AVAILABLE:
        st.error("PDF generation requires reportlab. Install with: pip install reportlab")
        return None
    
    try:
        pdf_io = BytesIO()
        doc = SimpleDocTemplate(
            pdf_io, 
            pagesize=letter,
            rightMargin=0.5*inch, 
            leftMargin=0.5*inch,
            topMargin=0.4*inch, 
            bottomMargin=0.4*inch
        )
        
        elements = []
        
        PRIMARY_COLOR = HexColor('#2B5797')
        SECONDARY_COLOR = HexColor('#505050')
        ACCENT_COLOR = HexColor('#0078D4')
        
        styles = getSampleStyleSheet()
        
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=PRIMARY_COLOR,
            spaceAfter=4,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=SECONDARY_COLOR,
            spaceAfter=4,
            alignment=TA_JUSTIFY
        )
        
        section_header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=11,
            textColor=PRIMARY_COLOR,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold'
        )
        
        header = resume_data.get('header', {})
        
        if header.get('name'):
            elements.append(Paragraph(header['name'].upper(), name_style))
        
        contact_items = []
        if header.get('email'):
            contact_items.append(header['email'])
        if header.get('phone'):
            contact_items.append(header['phone'])
        if header.get('location'):
            contact_items.append(header['location'])
        
        if contact_items:
            contact_style = ParagraphStyle('ContactStyle', parent=styles['Normal'], fontSize=9, textColor=SECONDARY_COLOR, alignment=TA_CENTER)
            elements.append(Paragraph('  •  '.join(contact_items), contact_style))
        
        elements.append(Spacer(1, 0.1*inch))
        elements.append(HRFlowable(width="100%", thickness=2, color=PRIMARY_COLOR, spaceAfter=0.1*inch))
        
        if resume_data.get('summary'):
            elements.append(Paragraph('PROFESSIONAL SUMMARY', section_header_style))
            elements.append(Paragraph(resume_data['summary'], body_style))
        
        skills = resume_data.get('skills_highlighted', [])
        if skills:
            elements.append(Paragraph('KEY SKILLS', section_header_style))
            skills_text = '  |  '.join(skills)
            skills_style = ParagraphStyle('SkillsStyle', parent=styles['Normal'], fontSize=9, textColor=PRIMARY_COLOR, alignment=TA_CENTER)
            elements.append(Paragraph(skills_text, skills_style))
        
        experience = resume_data.get('experience', [])
        if experience:
            elements.append(Paragraph('PROFESSIONAL EXPERIENCE', section_header_style))
            
            for exp in experience:
                if exp.get('title'):
                    job_style = ParagraphStyle('JobStyle', parent=styles['Normal'], fontSize=11, textColor=black, fontName='Helvetica-Bold')
                    elements.append(Paragraph(exp['title'], job_style))
                
                if exp.get('company') or exp.get('dates'):
                    company_style = ParagraphStyle('CompanyStyle', parent=styles['Normal'], fontSize=10, textColor=ACCENT_COLOR)
                    company_text = f"{exp.get('company', '')}  |  {exp.get('dates', '')}"
                    elements.append(Paragraph(company_text, company_style))
                
                bullets = exp.get('bullets', [])
                for bullet in bullets:
                    if bullet and bullet.strip():
                        bullet_style = ParagraphStyle('BulletStyle', parent=styles['Normal'], fontSize=10, textColor=SECONDARY_COLOR, leftIndent=15)
                        elements.append(Paragraph(f"▸  {bullet.strip()}", bullet_style))
                
                elements.append(Spacer(1, 0.1*inch))
        
        if resume_data.get('education'):
            elements.append(Paragraph('EDUCATION', section_header_style))
            elements.append(Paragraph(resume_data['education'], body_style))
        
        if resume_data.get('certifications'):
            elements.append(Paragraph('CERTIFICATIONS & ACHIEVEMENTS', section_header_style))
            elements.append(Paragraph(resume_data['certifications'], body_style))
        
        doc.build(elements)
        pdf_io.seek(0)
        return pdf_io
        
    except Exception as e:
        st.error(f"Error generating PDF: {e}")
        return None


def format_resume_as_text(resume_data):
    """Format structured resume JSON as plain text (from CareerLens)"""
    text = []
    
    header = resume_data.get('header', {})
    
    if header.get('name'):
        name = header['name'].upper()
        text.append("=" * 60)
        text.append(name.center(60))
        text.append("=" * 60)
        text.append("")
    
    if header.get('title'):
        text.append(header['title'].center(60))
        text.append("")
    
    contact = []
    if header.get('email'):
        contact.append(header['email'])
    if header.get('phone'):
        contact.append(header['phone'])
    if header.get('location'):
        contact.append(header['location'])
    
    if contact:
        text.append(' | '.join(contact))
        text.append("")
        text.append("-" * 60)
        text.append("")
    
    if resume_data.get('summary'):
        text.append("PROFESSIONAL SUMMARY")
        text.append("-" * 25)
        text.append(resume_data['summary'])
        text.append("")
    
    skills = resume_data.get('skills_highlighted', [])
    if skills:
        text.append("KEY SKILLS")
        text.append("-" * 25)
        for i in range(0, len(skills), 4):
            row_skills = skills[i:i+4]
            text.append("  •  ".join(row_skills))
        text.append("")
    
    experience = resume_data.get('experience', [])
    if experience:
        text.append("PROFESSIONAL EXPERIENCE")
        text.append("-" * 25)
        for exp in experience:
            job_line = ""
            if exp.get('title'):
                job_line = exp['title']
            if exp.get('company'):
                job_line += f" | {exp['company']}"
            if job_line:
                text.append(job_line)
            
            if exp.get('dates'):
                text.append(f"    {exp['dates']}")
            
            bullets = exp.get('bullets', [])
            for bullet in bullets:
                if bullet and bullet.strip():
                    text.append(f"    ▸ {bullet.strip()}")
            text.append("")
    
    if resume_data.get('education'):
        text.append("EDUCATION")
        text.append("-" * 25)
        text.append(resume_data['education'])
        text.append("")
    
    if resume_data.get('certifications'):
        text.append("CERTIFICATIONS & ACHIEVEMENTS")
        text.append("-" * 25)
        text.append(resume_data['certifications'])
    
    return '\n'.join(text)


# ============================================================================
# INDEED JOB SCRAPER (from CareerLens)
# ============================================================================

class IndeedScraperAPI:
    """Job scraper using Indeed Scraper API via RapidAPI (from CareerLens)"""
    def __init__(self, api_key):
        self.api_key = api_key
        self.url = "https://indeed-scraper-api.p.rapidapi.com/api/job"
        self.headers = {
            'Content-Type': 'application/json',
            'x-rapidapi-host': 'indeed-scraper-api.p.rapidapi.com',
            'x-rapidapi-key': api_key
        }
        self.rate_limiter = RateLimiter(max_requests_per_minute=3)
    
    def search_jobs(self, query, location="Hong Kong", max_rows=15, job_type="fulltime", country="hk"):
        """Search for jobs using Indeed Scraper API"""
        payload = {
            "scraper": {
                "maxRows": max_rows,
                "query": query,
                "location": location,
                "jobType": job_type,
                "radius": "50",
                "sort": "relevance",
                "fromDays": "7",
                "country": country
            }
        }
        
        try:
            print(f"🔍 Searching Indeed for '{query}' in {location}...")
            self.rate_limiter.wait_if_needed()
            
            def make_request():
                return requests.post(self.url, headers=self.headers, json=payload, timeout=60)
            
            response = api_call_with_retry(make_request, max_retries=3, initial_delay=3)
            
            if response and response.status_code == 201:
                data = response.json()
                jobs = []
                
                if 'returnvalue' in data and 'data' in data['returnvalue']:
                    job_list = data['returnvalue']['data']
                    
                    for job_data in job_list:
                        parsed_job = self._parse_job(job_data)
                        if parsed_job:
                            jobs.append(parsed_job)
                
                print(f"✅ Found {len(jobs)} jobs from Indeed")
                return jobs
            else:
                if response:
                    if response.status_code == 429:
                        print("❌ Rate limit reached for Indeed API")
                    else:
                        print(f"❌ Indeed API Error: {response.status_code}")
                return []
                
        except Exception as e:
            print(f"❌ Indeed search error: {e}")
            return []
    
    def _parse_job(self, job_data):
        """Parse job data from API response"""
        try:
            location_data = job_data.get('location', {})
            location = location_data.get('formattedAddressShort') or location_data.get('city', 'Hong Kong')
            
            job_types = job_data.get('jobType', [])
            job_type = ', '.join(job_types) if job_types else 'Full-time'
            
            benefits = job_data.get('benefits', [])
            attributes = job_data.get('attributes', [])
            
            full_description = job_data.get('descriptionText', 'No description')
            description = full_description[:50000] if len(full_description) > 50000 else full_description
            
            return {
                'id': hashlib.md5(f"{job_data.get('title', '')}_{job_data.get('companyName', '')}".encode()).hexdigest()[:12],
                'title': job_data.get('title', 'N/A'),
                'company': job_data.get('companyName', 'N/A'),
                'location': location,
                'description': description,
                'salary': 'Not specified',
                'job_type': job_type,
                'url': job_data.get('jobUrl', '#'),
                'posted_date': job_data.get('age', 'Recently'),
                'benefits': benefits[:5],
                'skills': attributes[:10],
                'company_rating': job_data.get('rating', {}).get('rating', 0),
                'is_remote': job_data.get('isRemote', False)
            }
        except Exception:
            return None


# ============================================================================
# RESUME PARSER - NO HARDCODED SKILLS
# ============================================================================

class ResumeParser:
    """Parse resume from PDF, DOCX, or TXT - Let GPT-4 extract skills"""
    
    def __init__(self):
        pass
    
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text from PDF file object"""
        try:
            PyPDF2 = _get_pypdf2()
            # Reset file position to beginning
            pdf_file.seek(0)
            text = ""
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text from DOCX file object"""
        try:
            from docx import Document  # Lazy load when needed
            # Reset file position to beginning
            docx_file.seek(0)
            doc = Document(docx_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def extract_text_from_txt(self, txt_file) -> str:
        """Extract text from TXT file object (from CareerLens)"""
        try:
            txt_file.seek(0)
            text = str(txt_file.read(), "utf-8")
            return text
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def extract_text(self, file_obj, filename: str) -> str:
        """Extract text from uploaded file (extended to support TXT)"""
        if filename.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_obj)
        elif filename.lower().endswith('.docx'):
            return self.extract_text_from_docx(file_obj)
        elif filename.lower().endswith('.txt'):
            return self.extract_text_from_txt(file_obj)
        else:
            raise ValueError("Unsupported file format. Use PDF, DOCX, or TXT.")
    
    def parse_resume(self, file_obj, filename: str) -> Dict:
        """Parse resume and extract raw text only"""
        try:
            text = self.extract_text(file_obj, filename)
            
            if not text or len(text.strip()) < 50:
                file_type = filename.split('.')[-1].upper() if '.' in filename else 'file'
                raise ValueError(
                    f"Could not extract sufficient text from resume. "
                    f"This may happen if:\n"
                    f"• The {file_type} is scanned/image-based (try a text-based document)\n"
                    f"• The file is corrupted or password-protected\n"
                    f"• The document is mostly empty\n"
                    f"Please try uploading a different format (PDF or DOCX with selectable text)."
                )
            
            resume_data = {
                'raw_text': text,
                'text_length': len(text),
                'word_count': len(text.split()),
                'filename': filename
            }
            
            return resume_data
            
        except Exception as e:
            raise Exception(f"Error parsing resume: {str(e)}")


# ============================================================================
# ENHANCED PROFILE EXTRACTION (from CareerLens)
# ============================================================================

def extract_relevant_resume_sections(resume_text):
    """Extract Experience and Education sections from resume to reduce token usage (from CareerLens)"""
    if not resume_text:
        return ""
    
    experience_keywords = [
        r'experience', r'work experience', r'employment', r'employment history',
        r'professional experience', r'work history', r'career history', r'positions held'
    ]
    education_keywords = [
        r'education', r'academic background', r'academic qualifications',
        r'educational background', r'qualifications', r'degrees'
    ]
    
    lines = resume_text.split('\n')
    relevant_sections = []
    current_section = None
    in_experience = False
    in_education = False
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        line_lower = line_stripped.lower()
        
        if any(re.search(rf'\b{kw}\b', line_lower) for kw in experience_keywords):
            if not in_experience:
                in_experience = True
                in_education = False
                if current_section:
                    relevant_sections.append(current_section)
                current_section = line + '\n'
            continue
        
        if any(re.search(rf'\b{kw}\b', line_lower) for kw in education_keywords):
            if not in_education:
                in_education = True
                if current_section:
                    relevant_sections.append(current_section)
                current_section = line + '\n'
            continue
        
        major_sections = [r'summary', r'objective', r'skills', r'certifications', 
                         r'awards', r'publications', r'projects', r'contact', r'personal']
        if any(re.search(rf'\b{section}\b', line_lower) for section in major_sections):
            if in_experience or in_education:
                if current_section:
                    relevant_sections.append(current_section)
                current_section = None
                in_experience = False
                in_education = False
            continue
        
        if in_experience or in_education:
            if current_section:
                current_section += line + '\n'
    
    if current_section and (in_experience or in_education):
        relevant_sections.append(current_section)
    
    result = '\n'.join(relevant_sections)
    
    if not result or len(result) < 100:
        date_pattern = r'\b(19|20)\d{2}\b|\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}'
        result_lines = []
        for line in lines:
            if re.search(date_pattern, line, re.IGNORECASE):
                result_lines.append(line)
            elif result_lines:
                if len([l for l in result_lines[-3:] if l.strip()]) < 3:
                    result_lines.append(line)
                else:
                    break
        if result_lines:
            result = '\n'.join(result_lines[:50])
    
    if result:
        return result[:2000] if len(result) > 2000 else result
    
    return ""


def extract_structured_profile(resume_text, enable_verification=False):
    """Extract structured profile from resume with optional two-pass verification (from CareerLens)"""
    try:
        # Check if API keys are configured
        is_configured, error_msg = Config.check_azure_credentials()
        if not is_configured:
            print(f"❌ Configuration Error: {error_msg}")
            return None
        
        client = AzureOpenAI(
            azure_endpoint=Config.AZURE_ENDPOINT,
            api_key=Config.AZURE_API_KEY,
            api_version=Config.AZURE_API_VERSION
        )
        
        # FIRST PASS: Initial extraction
        prompt_pass1 = f"""You are an expert at parsing resumes. Extract structured information from the following resume text.

RESUME TEXT:
{resume_text[:6000]}

Please extract and return the following information in JSON format:
{{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL if mentioned",
    "portfolio": "Portfolio/website URL if mentioned",
    "summary": "Professional summary or objective (2-3 sentences)",
    "experience": "Work experience with job titles, companies, dates, and achievements",
    "education": "Education details including degrees, institutions, and graduation dates",
    "skills": "Comma-separated list of technical and soft skills",
    "certifications": "Professional certifications, awards, or achievements"
}}

Important:
- If information is not found, use "N/A" or empty string
- Extract all relevant skills mentioned
- Keep the summary concise but informative
- Return ONLY valid JSON, no additional text"""
        
        print("🤖 Pass 1: Extracting profile information...")
        response_pass1 = client.chat.completions.create(
            model=Config.AZURE_MODEL,
            messages=[
                {"role": "system", "content": "You are a resume parser. Extract structured information and return only valid JSON."},
                {"role": "user", "content": prompt_pass1}
            ],
            max_tokens=2000,
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        
        content_pass1 = response_pass1.choices[0].message.content
        profile_data_pass1 = json.loads(content_pass1)
        
        if not enable_verification:
            print("✅ Profile extraction complete (single pass)")
            return profile_data_pass1
        
        # SECOND PASS: Self-correction (optional)
        relevant_sections = extract_relevant_resume_sections(resume_text)
        
        if relevant_sections:
            resume_context = f"""RELEVANT RESUME SECTIONS (Experience and Education only):
{relevant_sections}"""
        else:
            resume_context = f"""RELEVANT RESUME SECTIONS (limited):
{resume_text[:1500]}"""
        
        prompt_pass2 = f"""You are a resume quality checker. Review the extracted profile data against the relevant resume sections and verify accuracy, especially for dates and company names.

{resume_context}

EXTRACTED PROFILE DATA (from first pass):
{json.dumps(profile_data_pass1, indent=2)}

Please review and correct the extracted data, paying special attention to:
1. **Dates** - Verify all employment dates, education dates, and certification dates are accurate
2. **Company Names** - Verify all company/organization names are spelled correctly
3. **Job Titles** - Verify job titles are accurate
4. **Education Institutions** - Verify institution names are correct

Return the corrected profile data in the same JSON format. If everything is correct, return the data as-is.

Return ONLY valid JSON, no additional text."""
        
        print("🔍 Pass 2: Verifying profile data...")
        response_pass2 = client.chat.completions.create(
            model=Config.AZURE_MODEL,
            messages=[
                {"role": "system", "content": "You are a resume quality checker. Verify and correct extracted data. Return only valid JSON."},
                {"role": "user", "content": prompt_pass2}
            ],
            max_tokens=2000,
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        content_pass2 = response_pass2.choices[0].message.content
        profile_data_corrected = json.loads(content_pass2)
        print("✅ Profile extraction complete (two-pass verification)")
        return profile_data_corrected
        
    except Exception as e:
        print(f"❌ Profile extraction error: {e}")
        return None


def generate_tailored_resume(user_profile, job_posting, raw_resume_text=None):
    """Generate a tailored resume based on user profile and job posting (from CareerLens)"""
    try:
        # Check if API keys are configured
        is_configured, error_msg = Config.check_azure_credentials()
        if not is_configured:
            print(f"❌ Configuration Error: {error_msg}")
            return None
        
        client = AzureOpenAI(
            azure_endpoint=Config.AZURE_ENDPOINT,
            api_key=Config.AZURE_API_KEY,
            api_version=Config.AZURE_API_VERSION
        )
        
        system_instructions = """You are an expert resume writer with expertise in ATS optimization and career coaching.
Your task is to create a tailored resume by analyzing the job description and adapting the user's profile.
Return ONLY valid JSON - no markdown, no additional text, no code blocks."""

        job_description = f"""JOB POSTING TO MATCH:
Title: {job_posting.get('title', 'N/A')}
Company: {job_posting.get('company', 'N/A')}
Description: {job_posting.get('description', 'N/A')[:3000]}
Required Skills: {', '.join(job_posting.get('skills', [])[:10]) if job_posting.get('skills') else 'N/A'}"""

        structured_profile = f"""STRUCTURED PROFILE:
Name: {user_profile.get('name', 'N/A')}
Email: {user_profile.get('email', 'N/A')}
Phone: {user_profile.get('phone', 'N/A')}
Location: {user_profile.get('location', 'N/A')}
LinkedIn: {user_profile.get('linkedin', 'N/A')}
Summary: {user_profile.get('summary', 'N/A')}
Experience: {user_profile.get('experience', 'N/A')[:2000]}
Education: {user_profile.get('education', 'N/A')}
Skills: {user_profile.get('skills', 'N/A')}
Certifications: {user_profile.get('certifications', 'N/A')}"""

        raw_resume_section = ""
        if raw_resume_text:
            raw_resume_section = f"\n\nORIGINAL RESUME TEXT (for reference):\n{raw_resume_text[:2000]}"

        prompt = f"""{system_instructions}

{job_description}

{structured_profile}{raw_resume_section}

INSTRUCTIONS:
1. Analyze the job posting and identify key skills, technologies, and qualifications needed
2. Tailor the profile to match by:
   - Rewriting the summary to emphasize relevant experience
   - Highlighting skills that match job requirements
   - Rewriting experience bullet points to emphasize relevant achievements
   - Using keywords from the job description for ATS optimization
3. Focus on achievements and measurable results
4. Maintain accuracy - only use information from the provided profile

Return your response as a JSON object with this structure:
{{
  "header": {{
    "name": "Full Name",
    "title": "Professional Title (tailored to job)",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL or empty string"
  }},
  "summary": "2-3 sentence professional summary tailored to the job",
  "skills_highlighted": ["Skill 1", "Skill 2", "Skill 3", ...],
  "experience": [
    {{
      "company": "Company Name",
      "title": "Job Title",
      "dates": "Date Range",
      "bullets": ["Achievement bullet 1...", "Achievement bullet 2..."]
    }}
  ],
  "education": "Education details",
  "certifications": "Certifications and achievements"
}}

Return ONLY the JSON object."""
        
        print("✨ Generating tailored resume...")
        response = client.chat.completions.create(
            model=Config.AZURE_MODEL,
            messages=[
                {"role": "system", "content": system_instructions},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000,
            temperature=0.7,
            response_format={"type": "json_object"}
        )
        
        content = response.choices[0].message.content
        resume_data = json.loads(content)
        print("✅ Tailored resume generated!")
        return resume_data
        
    except Exception as e:
        print(f"❌ Resume generation error: {e}")
        return None


# ============================================================================
# GPT-4 JOB ROLE DETECTOR - EXTRACTS SKILLS DYNAMICALLY
# ============================================================================

class GPT4JobRoleDetector:
    """Use GPT-4 to detect job roles AND extract skills dynamically"""
    
    def __init__(self):
        self._client = None
        self.model = Config.AZURE_MODEL
    
    @property
    def client(self):
        """Lazy-load AzureOpenAI client only when needed"""
        if self._client is None:
            self._client = AzureOpenAI(
                azure_endpoint=Config.AZURE_ENDPOINT,
                api_key=Config.AZURE_API_KEY,
                api_version=Config.AZURE_API_VERSION
            )
        return self._client
    
    def analyze_resume_for_job_roles(self, resume_data: Dict) -> Dict:
        """Analyze resume with GPT-4 - Extract ALL skills dynamically"""
        
        resume_text = resume_data.get('raw_text', '')[:3000]
        
        system_prompt = """You are an expert career advisor and resume analyst.

Analyze the resume and extract:
1. ALL skills (technical, soft skills, tools, languages, frameworks, methodologies, domain knowledge)
2. Job role recommendations
3. Seniority level
4. SIMPLE job search keywords (for job board APIs)

IMPORTANT for job search:
- Provide a SIMPLE primary role (e.g., "Program Manager" not complex OR/AND queries)
- Keep search keywords SHORT and COMMON
- Avoid complex boolean logic in search queries

Return JSON with this EXACT structure:
{
    "primary_role": "Simple job title (e.g., Program Manager)",
    "simple_search_terms": ["term1", "term2", "term3"],
    "confidence": 0.95,
    "seniority_level": "Junior/Mid-Level/Senior/Lead/Executive",
    "skills": ["skill1", "skill2", "skill3", ...],
    "core_strengths": ["strength1", "strength2", "strength3"],
    "job_search_keywords": ["keyword1", "keyword2"],
    "optimal_search_query": "Simple search string (just the job title)",
    "location_preference": "Detected or 'United States'",
    "industries": ["industry1", "industry2"],
    "alternative_roles": ["role1", "role2", "role3"]
}"""

        user_prompt = f"""Analyze this resume and extract ALL information:

RESUME:
{resume_text}

IMPORTANT - Extract ALL skills including:
- Programming languages (Python, R, SQL, etc.)
- Tools and software (Tableau, Salesforce, Excel, etc.)
- Methodologies (Agile, Scrum, Kanban, etc.)
- Soft skills (Leadership, Communication, etc.)
- Domain expertise (Banking, Finance, Analytics, etc.)
- Technical skills (Data Analysis, Machine Learning, etc.)
- Languages (English, Cantonese, Mandarin, etc.)

For job search, provide SIMPLE terms that would work on LinkedIn/Indeed (not complex boolean queries).

Be thorough and creative!"""

        try:
            # Check if API keys are configured before attempting API call
            is_configured, error_msg = Config.check_azure_credentials()
            if not is_configured:
                print(f"❌ Configuration Error: {error_msg}")
                fallback = self._fallback_analysis()
                fallback['_error'] = error_msg
                fallback['_analysis_failed'] = True
                return fallback
            
            print("🤖 Calling GPT-4 for resume analysis...")
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )
            
            ai_analysis = json.loads(response.choices[0].message.content)
            print(f"✅ GPT-4 analysis complete! Found {len(ai_analysis.get('skills', []))} skills")
            
            # Validate that we got meaningful data - if primary_role is empty or generic, flag it
            if not ai_analysis.get('primary_role') or ai_analysis.get('primary_role') == 'Professional':
                ai_analysis['_analysis_incomplete'] = True
            
            return ai_analysis
            
        except Exception as e:
            error_msg = str(e)
            print(f"❌ GPT-4 Error: {error_msg}")
            # Store error for UI to display
            fallback = self._fallback_analysis()
            fallback['_error'] = error_msg
            fallback['_analysis_failed'] = True
            return fallback
    
    def _fallback_analysis(self) -> Dict:
        """Fallback if GPT-4 fails - returns empty strings so user can fill in manually"""
        return {
            "primary_role": "",  # Empty so user fills in
            "simple_search_terms": [],
            "confidence": 0.0,
            "seniority_level": "",
            "skills": [],
            "core_strengths": [],
            "job_search_keywords": [],
            "optimal_search_query": "",
            "location_preference": "",
            "industries": [],
            "alternative_roles": [],
            "_analysis_failed": True  # Flag to indicate fallback was used
        }


# ============================================================================
# LINKEDIN JOB SEARCHER - WITH BETTER ERROR HANDLING
# ============================================================================

class LinkedInJobSearcher:
    """Search for jobs using RapidAPI LinkedIn API"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://linkedin-job-search-api.p.rapidapi.com/active-jb-7d"
        self.headers = {
            "x-rapidapi-key": api_key,
            "x-rapidapi-host": "linkedin-job-search-api.p.rapidapi.com"
        }
    
    def test_api_connection(self) -> Tuple[bool, str]:
        """Test if the API is working"""
        try:
            querystring = {
                "limit": "5",
                "offset": "0",
                "title_filter": "\"Engineer\"",
                "location_filter": "\"Hong Kong\"",
                "description_type": "text"
            }
            
            response = requests.get(
                self.base_url,
                headers=self.headers,
                params=querystring,
                timeout=10
            )
            
            if response.status_code == 200:
                return True, "API is working"
            elif response.status_code == 403:
                return False, "API key is invalid or expired (403 Forbidden)"
            elif response.status_code == 429:
                return False, "Rate limit exceeded (429 Too Many Requests)"
            else:
                return False, f"API returned status code {response.status_code}"
        
        except Exception as e:
            return False, f"Connection error: {str(e)}"
    
    def search_jobs(
        self,
        keywords: str,
        location: str = "Hong Kong",
        limit: int = 20
    ) -> List[Dict]:
        """Search LinkedIn jobs with simplified queries"""
        
        # Simplify complex queries
        simple_keywords = self._simplify_query(keywords)
        
        querystring = {
            "limit": str(limit),
            "offset": "0",
            "title_filter": f'"{simple_keywords}"',
            "location_filter": f'"{location}"',
            "description_type": "text"
        }
        
        try:
            print(f"🔍 Searching RapidAPI...")
            print(f"   Original query: {keywords}")
            print(f"   Simplified to: {simple_keywords}")
            print(f"   Location: {location}")
            
            response = requests.get(
                self.base_url, 
                headers=self.headers, 
                params=querystring, 
                timeout=30
            )
            
            print(f"📊 API Response Status: {response.status_code}")
            
            if response.status_code == 403:
                print("❌ API Key Error: 403 Forbidden")
                print("   Your RapidAPI key might be invalid or expired")
                print("   Check: https://rapidapi.com/")
                return []
            
            elif response.status_code == 429:
                print("❌ Rate Limit: 429 Too Many Requests")
                print("   Wait a few minutes or upgrade your RapidAPI plan")
                return []
            
            elif response.status_code != 200:
                print(f"❌ API Error: {response.status_code}")
                print(f"   Response: {response.text[:200]}")
                return []
            
            data = response.json()
            
            # Handle different response formats
            if isinstance(data, list):
                jobs = data
            elif isinstance(data, dict):
                jobs = data.get('data', data.get('jobs', data.get('results', [])))
            else:
                jobs = []
            
            if not jobs:
                print(f"⚠️ No jobs found for '{simple_keywords}'")
                print("   Trying fallback searches...")
                
                # Try alternative searches
                for alternative in self._get_alternative_searches(simple_keywords):
                    alt_jobs = self._try_alternative_search(alternative, location, 10)
                    if alt_jobs:
                        print(f"✅ Found {len(alt_jobs)} jobs with alternative search: {alternative}")
                        jobs.extend(alt_jobs)
                        if len(jobs) >= 10:
                            break
            
            normalized = self._normalize_jobs(jobs)
            print(f"✅ Retrieved {len(normalized)} jobs from RapidAPI")
            return normalized
            
        except Exception as e:
            print(f"❌ LinkedIn API Error: {str(e)}")
            return []
    
    def _simplify_query(self, query: str) -> str:
        """Simplify complex boolean queries to simple terms"""
        # Remove boolean operators and parentheses
        simple = query.replace(" OR ", " ").replace(" AND ", " ")
        simple = simple.replace("(", "").replace(")", "")
        simple = simple.replace('"', "")
        
        # Take first few words (most important)
        words = simple.split()[:3]
        return " ".join(words)
    
    def _get_alternative_searches(self, primary_query: str) -> List[str]:
        """Generate alternative search terms"""
        alternatives = [
            primary_query.split()[0] if primary_query.split() else primary_query,  # First word only
            "Manager",  # Generic fallback
            "Analyst",  # Generic fallback
        ]
        return alternatives
    
    def _try_alternative_search(self, keywords: str, location: str, limit: int) -> List[Dict]:
        """Try an alternative search"""
        try:
            querystring = {
                "limit": str(limit),
                "offset": "0",
                "title_filter": f'"{keywords}"',
                "location_filter": f'"{location}"',
                "description_type": "text"
            }
            
            response = requests.get(
                self.base_url,
                headers=self.headers,
                params=querystring,
                timeout=20
            )
            
            if response.status_code == 200:
                data = response.json()
                if isinstance(data, list):
                    return data
                elif isinstance(data, dict):
                    return data.get('data', data.get('jobs', data.get('results', [])))
            
            return []
        
        except Exception:
            return []
    
    def _normalize_jobs(self, jobs: List[Dict]) -> List[Dict]:
        """Normalize job structure"""
        normalized_jobs = []
        
        for job in jobs:
            try:
                # Handle location
                location = "Remote"
                if job.get('locations_derived') and len(job['locations_derived']) > 0:
                    location = job['locations_derived'][0]
                elif job.get('locations_raw'):
                    try:
                        loc_raw = job['locations_raw'][0]
                        if isinstance(loc_raw, dict) and 'address' in loc_raw:
                            addr = loc_raw['address']
                            city = addr.get('addressLocality', '')
                            region = addr.get('addressRegion', '')
                            if city and region:
                                location = f"{city}, {region}"
                    except (KeyError, TypeError, IndexError):
                        pass
                
                normalized_job = {
                    'id': job.get('id', f"job_{len(normalized_jobs)}"),
                    'title': job.get('title', 'Unknown Title'),
                    'company': job.get('organization', 'Unknown Company'),
                    'location': location,
                    'description': job.get('description_text', ''),
                    'url': job.get('url', ''),
                    'posted_date': job.get('date_posted', 'Unknown'),
                }
                
                normalized_jobs.append(normalized_job)
                
            except Exception as e:
                continue
        
        return normalized_jobs


# ============================================================================
# CACHED MODEL LOADING - Prevents re-downloading on every page load
# ============================================================================

@st.cache_resource(show_spinner=False)
def _get_sentence_transformer_model_cached():
    """Load and cache SentenceTransformer model - only loaded once"""
    print("📦 Loading sentence transformer model (first time only)...")
    SentenceTransformer = _get_sentence_transformer_class()
    model = SentenceTransformer(Config.MODEL_NAME)
    print("✅ Model loaded and cached!")
    return model


@st.cache_resource(show_spinner=False)
def _get_pinecone_client_cached():
    """Get cached Pinecone client - only initialized once"""
    print("🔗 Initializing Pinecone client (first time only)...")
    Pinecone, _ = _get_pinecone_classes()
    pc = Pinecone(api_key=Config.PINECONE_API_KEY)
    print("✅ Pinecone client cached!")
    return pc


@st.cache_resource(show_spinner=False)
def _get_pinecone_index_cached(_pc):
    """Get cached Pinecone index - only initialized once"""
    _, ServerlessSpec = _get_pinecone_classes()
    existing_indexes = _pc.list_indexes()
    index_names = [idx['name'] for idx in existing_indexes]
    
    if Config.INDEX_NAME not in index_names:
        print(f"🔨 Creating new Pinecone index: {Config.INDEX_NAME}")
        _pc.create_index(
            name=Config.INDEX_NAME,
            dimension=Config.EMBEDDING_DIMENSION,
            metric='cosine',
            spec=ServerlessSpec(
                cloud='aws',
                region=Config.PINECONE_ENVIRONMENT
            )
        )
        time.sleep(2)
    else:
        print(f"✅ Using existing Pinecone index: {Config.INDEX_NAME}")
    
    return _pc.Index(Config.INDEX_NAME)


# ============================================================================
# JOB MATCHER - PINECONE SEMANTIC SEARCH & RANKING
# ============================================================================

class JobMatcher:
    """Match resume to jobs using Pinecone semantic search and skill matching"""
    
    def __init__(self):
        # Lazy initialization - resources are loaded only when first accessed via properties
        # This avoids loading ~85MB model on every rerun
        self._pc = None
        self._model = None
        self._index = None
    
    @property
    def pc(self):
        """Lazy-load Pinecone client"""
        if self._pc is None:
            self._pc = _get_pinecone_client_cached()
        return self._pc
    
    @property
    def model(self):
        """Lazy-load SentenceTransformer model"""
        if self._model is None:
            self._model = _get_sentence_transformer_model_cached()
        return self._model
    
    @property
    def index(self):
        """Lazy-load Pinecone index"""
        if self._index is None:
            self._index = _get_pinecone_index_cached(self.pc)
        return self._index
    
    # _initialize_index is no longer needed - index is lazy-loaded via property
    
    def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding vector"""
        text = str(text).strip()
        if not text:
            text = "empty"
        
        embedding = self.model.encode(text, convert_to_tensor=False)
        return embedding.tolist()
    
    def index_jobs(self, jobs: List[Dict]) -> int:
        """Index jobs in Pinecone"""
        if not jobs:
            return 0
        
        vectors_to_upsert = []
        
        for job in jobs:
            try:
                job_text = f"{job['title']} {job['company']} {job['description']}"
                embedding = self.generate_embedding(job_text)
                
                vectors_to_upsert.append({
                    'id': job['id'],
                    'values': embedding,
                    'metadata': {
                        'title': job['title'][:512],
                        'company': job['company'][:512],
                        'location': job['location'][:512],
                        'description': job['description'][:1000],
                        'url': job.get('url', '')[:512],
                        'posted_date': str(job.get('posted_date', ''))[:100]
                    }
                })
                
            except Exception as e:
                print(f"⚠️ Error indexing job {job.get('id', 'unknown')}: {e}")
                continue
        
        if vectors_to_upsert:
            self.index.upsert(vectors=vectors_to_upsert)
            return len(vectors_to_upsert)
        
        return 0
    
    def search_similar_jobs(self, resume_data: Dict, ai_analysis: Dict, top_k: int = 20) -> List[Dict]:
        """Search for similar jobs using semantic similarity"""
        try:
            # Create rich query from resume + AI analysis
            primary_role = ai_analysis.get('primary_role', '')
            skills = ' '.join(ai_analysis.get('skills', [])[:20])
            resume_snippet = resume_data.get('raw_text', '')[:1000]
            
            query_text = f"{primary_role} {skills} {resume_snippet}"
            
            print(f"🎯 Creating semantic embedding for resume...")
            query_embedding = self.generate_embedding(query_text)
            
            print(f"🔍 Searching Pinecone for top {top_k} matches...")
            results = self.index.query(
                vector=query_embedding,
                top_k=top_k,
                include_metadata=True
            )
            
            matched_jobs = []
            for match in results['matches']:
                job = {
                    'id': match['id'],
                    'similarity_score': float(match['score']) * 100,
                    **match['metadata']
                }
                matched_jobs.append(job)
            
            print(f"✅ Found {len(matched_jobs)} semantic matches")
            return matched_jobs
            
        except Exception as e:
            print(f"❌ Search error: {e}")
            return []


# ============================================================================
# MAIN BACKEND - ORCHESTRATES EVERYTHING
# ============================================================================

class JobSeekerBackend:
    """Main backend with FULL integration - optimized for fast startup"""
    
    def __init__(self):
        print("🚀 Initializing Job Matcher Backend (lightweight)...")
        Config.validate()
        
        # Lightweight components - instant init
        self.resume_parser = ResumeParser()
        self.gpt4_detector = GPT4JobRoleDetector()
        
        # Lazy-load heavy components - deferred until first use
        self._job_searcher = None
        self._matcher = None
        
        print("✅ Backend initialized (fast mode)!\n")
    
    @property
    def matcher(self):
        """Lazy-load JobMatcher only when needed"""
        if self._matcher is None:
            print("📦 Loading JobMatcher (first use)...")
            self._matcher = JobMatcher()
        return self._matcher
    
    @property
    def job_searcher(self):
        """Lazy initialization of job searcher - only tests connection when first used"""
        if self._job_searcher is None:
            print("\n🧪 Initializing RapidAPI job searcher...")
            
            # Check if RAPIDAPI_KEY is configured
            if not Config.RAPIDAPI_KEY:
                print("⚠️ WARNING: RAPIDAPI_KEY is not configured!")
                print("   Job search functionality will not work.")
                print("   Please configure RAPIDAPI_KEY in your Streamlit secrets.")
                # Return a placeholder that will fail gracefully
                self._job_searcher = LinkedInJobSearcher("")  # Empty key - will fail API calls
                return self._job_searcher
            
            self._job_searcher = LinkedInJobSearcher(Config.RAPIDAPI_KEY)
            # Test API connection only once
            is_working, message = self._job_searcher.test_api_connection()
            if is_working:
                print(f"✅ {message}")
            else:
                print(f"⚠️ WARNING: {message}")
                print("   Job search may not work properly!")
        return self._job_searcher
    
    def test_api_connection(self):
        """Test API connection on demand (not at startup)"""
        return self.job_searcher.test_api_connection()
    
    def process_resume(self, file_obj, filename: str) -> Tuple[Dict, Dict]:
        """Process resume and get AI analysis"""
        print(f"📄 Processing resume: {filename}")
        
        # Parse resume
        resume_data = self.resume_parser.parse_resume(file_obj, filename)
        print(f"✅ Extracted {resume_data['word_count']} words from resume")
        
        # Get GPT-4 analysis
        ai_analysis = self.gpt4_detector.analyze_resume_for_job_roles(resume_data)
        
        # Add skills to resume_data
        resume_data['skills'] = ai_analysis.get('skills', [])
        
        return resume_data, ai_analysis
    
    def search_and_match_jobs(self, resume_data: Dict, ai_analysis: Dict, num_jobs: int = 30, 
                               search_keywords: str = None, location: str = None) -> List[Dict]:
        """Search for jobs and rank by match quality.
        
        Args:
            resume_data: Parsed resume data
            ai_analysis: AI-extracted skills and role analysis
            num_jobs: Number of jobs to search for
            search_keywords: Search keywords (if None, uses ai_analysis primary_role)
            location: Location preference (if None, defaults to Hong Kong)
        """
        # Use provided keywords or fall back to AI-detected role
        # Avoid using generic "Professional" as search query - it returns poor results
        primary_role = ai_analysis.get('primary_role', '')
        search_query = search_keywords if search_keywords else primary_role
        
        # If no search query available, return empty to prompt user to provide keywords
        if not search_query or not search_query.strip():
            print("⚠️ No search keywords provided and no primary role detected.")
            print("   Please provide search keywords in your profile.")
            return []
        
        location = location if location else "Hong Kong"
        
        print(f"\n{'='*60}")
        print(f"🔍 SEARCHING JOBS")
        print(f"{'='*60}")
        print(f"🔍 Search Query: {search_query}")
        print(f"📍 Location: {location}")
        print(f"{'='*60}\n")
        
        # Search jobs
        jobs = self.job_searcher.search_jobs(
            keywords=search_query,
            location=location,
            limit=num_jobs
        )
        
        if not jobs or len(jobs) == 0:
            print("\n❌ No jobs found from RapidAPI")
            print("\n💡 Possible reasons:")
            print("   - API key might be invalid/expired")
            print("   - Rate limit exceeded")
            print("   - No jobs available for this search term")
            print("\n🔧 Suggestions:")
            print("   - Check your RapidAPI account at https://rapidapi.com/")
            print("   - Wait a few minutes if rate limited")
            print("   - Try with a different resume/role")
            return []
        
        print(f"\n✅ Retrieved {len(jobs)} jobs from RapidAPI")
        print(f"📊 Indexing jobs in Pinecone...")
        
        # Index jobs
        indexed = self.matcher.index_jobs(jobs)
        print(f"✅ Indexed {indexed} jobs in vector database")
        
        # Wait for indexing (reduced from 2s to 1s for faster response)
        print("⏳ Waiting for indexing to complete...")
        time.sleep(1)
        
        # Match resume to jobs
        print(f"\n🎯 MATCHING & RANKING JOBS")
        print(f"{'='*60}")
        matched_jobs = self.matcher.search_similar_jobs(
            resume_data, 
            ai_analysis, 
            top_k=min(20, len(jobs))
        )
        
        if not matched_jobs:
            print("⚠️ No matches found")
            return []
        
        # Calculate match scores
        matched_jobs = self._calculate_match_scores(matched_jobs, ai_analysis)
        
        # Sort by combined score
        matched_jobs.sort(key=lambda x: x.get('combined_score', 0), reverse=True)
        
        print(f"✅ Ranked {len(matched_jobs)} jobs by match quality")
        print(f"{'='*60}\n")
        
        return matched_jobs
    
    def _calculate_match_scores(self, jobs: List[Dict], ai_analysis: Dict) -> List[Dict]:
        """Calculate detailed match scores - 60% semantic + 40% skill match"""
        
        candidate_skills = set([s.lower() for s in ai_analysis.get('skills', [])])
        
        print(f"📊 Calculating match scores using {len(candidate_skills)} candidate skills...")
        
        for job in jobs:
            description = job.get('description', '').lower()
            title = job.get('title', '').lower()
            
            # Count skill matches
            matched_skills = []
            for skill in candidate_skills:
                if skill in description or skill in title:
                    matched_skills.append(skill)
            
            # Calculate skill match percentage
            skill_match_pct = (len(matched_skills) / len(candidate_skills) * 100) if candidate_skills else 0
            
            # Semantic similarity (from Pinecone)
            semantic_score = job.get('similarity_score', 0)
            
            # Combined score: 60% semantic + 40% skill match
            combined_score = (0.6 * semantic_score) + (0.4 * skill_match_pct)
            
            # Add to job
            job['skill_match_percentage'] = round(skill_match_pct, 1)
            job['matched_skills'] = list(matched_skills)[:10]
            job['matched_skills_count'] = len(matched_skills)
            job['combined_score'] = round(combined_score, 1)
            job['semantic_score'] = round(semantic_score, 1)
        
        return jobs
    
    @staticmethod
    def parse_cv_with_ai(cv_text):
        prompt = f"""
Below is the cv text of a candidate. 
Please extract structured information (leave blank if missing):
cv_text: '''{cv_text}'''

Please output JSON, fields including:
- education_level(doctor/master/bachelor/associate/highschool)
- major
- graduation_status(fresh graduate/experienced/in study)
- university_background(985 university/211 university/overseas university/regular university/other)
- languages
- certificates
- hard_skills
- soft_skills
- work_experience(fresh graduate/1-3 years/3-5 years/5-10 years/10+ years)
- project_experience
- location_preference
- industry_preference
- salary_expectation
- benefits_expectation

Please return the result in the JSON format only, no extra explanation.
"""

        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )

        try:
            return json.loads(response.choices[0].message.content)
        except Exception:
            return {}

class JobMatcherBackend:
    """Main backend with FULL integration"""
    
    def fetch_real_jobs(self, search_query, location="", country="us", num_pages=1):
        """Get actual job data from JSearch API"""
        try:
            # JSearch API configuration
            API_KEY = "your_jsearch_api_key_here"  # You need to get api key from https://jsearch.app/
            BASE_URL = "https://jsearch.p.rapidapi.com/search"
            
            headers = {
                "X-RapidAPI-Key": API_KEY,
                "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
            }
            
            all_jobs = []
            
            for page in range(1, num_pages + 1):
                querystring = {
                    "query": f"{search_query} {location}",
                    "page": str(page),
                    "num_pages": "1"
                }
                
                response = requests.get(BASE_URL, headers=headers, params=querystring)
                
                if response.status_code == 200:
                    data = response.json()
                    jobs = data.get('data', [])
                    all_jobs.extend(jobs)
                    print(f"✅ Page {page} fetched {len(jobs)} jobs")
                else:
                    print(f"❌ API request failed: {response.status_code}")
                    break
            
            print(f"🎯 Found total of {len(all_jobs)} positions")
            return all_jobs
            
        except Exception as e:
            print(f"❌ Failed to fetch jobs: {e}")
            # return simulated data as fallback
            return self.get_mock_jobs(search_query, location)

    def get_mock_jobs(self, search_query, location):
        """return mock job data (used when API is unavailable)"""
        print("🔄 Using simulated data...")
        
        mock_jobs = [
            {
                'job_title': f'Senior {search_query}',
                'employer_name': 'Tech Company Inc.',
                'job_city': location or 'Hong Kong',
                'job_country': 'HK',
                'job_employment_type': 'FULLTIME',
                'job_posted_at': '2024-01-15',
                'job_description': f'We are looking for a skilled {search_query} to join our team. Requirements include strong programming skills and experience.',
                'job_apply_link': 'https://example.com/apply/1',
                'job_highlights': {
                    'Qualifications': ['Bachelor\'s degree in Computer Science', '3+ years of experience'],
                    'Responsibilities': ['Develop software applications', 'Collaborate with team members']
                }
            },
            {
                'job_title': f'Junior {search_query}',
                'employer_name': 'Startup Solutions',
                'job_city': location or 'Hong Kong',
                'job_country': 'HK',
                'job_employment_type': 'FULLTIME',
                'job_posted_at': '2024-01-10',
                'job_description': f'Entry-level position for {search_query}. Great learning opportunity for recent graduates.',
                'job_apply_link': 'https://example.com/apply/2',
                'job_highlights': {
                    'Qualifications': ['Degree in related field', 'Basic programming knowledge'],
                    'Responsibilities': ['Assist senior developers', 'Learn new technologies']
                }
            },
            {
                'job_title': f'{search_query} Specialist',
                'employer_name': 'Global Corp',
                'job_city': location or 'Hong Kong',
                'job_country': 'HK',
                'job_employment_type': 'CONTRACTOR',
                'job_posted_at': '2024-01-08',
                'job_description': f'Contract position for {search_query} with potential for extension.',
                'job_apply_link': 'https://example.com/apply/3',
                'job_highlights': {
                    'Qualifications': ['Proven track record', 'Excellent communication skills'],
                    'Responsibilities': ['Project development', 'Client meetings']
                }
            }
        ]
        
        return mock_jobs

    def calculate_job_match_score(self, job_seeker_data, job_data):
        """calcalate job match score between job seeker and job data"""
        try:
            score = 0
            max_score = 100
            matched_skills = []
            
            # 1. Skill match (40%)
            job_seeker_skills = job_seeker_data.get('hard_skills', '').lower()
            job_description = job_data.get('job_description', '').lower()
            
            if job_seeker_skills:
                skills_list = [skill.strip().lower() for skill in job_seeker_skills.split(',')]
                for skill in skills_list:
                    if skill and skill in job_description:
                        score += 5  # Each match score add 5 points
                        matched_skills.append(skill)
                        if score >= 40:  # Max skill points at 40
                            score = 40
                            break
            
            # 2. Experience match (20%)
            job_seeker_experience = job_seeker_data.get('work_experience', '').lower()
            if 'senior' in job_data.get('job_title', '').lower() and 'senior' in job_seeker_experience.lower():
                score += 20
            elif 'junior' in job_data.get('job_title', '').lower() and 'junior' in job_seeker_experience.lower():
                score += 20
            elif 'entry' in job_data.get('job_title', '').lower() and 'fresh' in job_seeker_experience.lower():
                score += 20
            else:
                score += 10  # 10 points for general experience match
            
            # 3. Location match (20%)
            job_seeker_location = job_seeker_data.get('location_preference', '').lower()
            job_location = job_data.get('job_city', '').lower()
            
            if job_seeker_location and job_location:
                if job_seeker_location in job_location or job_location in job_seeker_location:
                    score += 20
                else:
                    score += 5 # Unmatched location but give base score of 5
            
            # 4. Job Title Match (20%)
            job_seeker_role = job_seeker_data.get('primary_role', '').lower()
            job_title = job_data.get('job_title', '').lower()
            
            if job_seeker_role and job_title:
                if job_seeker_role in job_title:
                    score += 20
                else:
                    # Searching for keywords in job title
                    search_terms = job_seeker_data.get('simple_search_terms', '').lower()
                    if search_terms:
                        terms = [term.strip() for term in search_terms.split(',')]
                        for term in terms:
                            if term in job_title:
                                score += 15
                                break
            
            # Make sure the score is between 0 and 100
            score = min(max(score, 0), 100)
            
            return {
                'overall_score': score,
                'matched_skills': matched_skills,
                'skill_match': len(matched_skills),
                'experience_match': 'senior' in job_seeker_experience and 'senior' in job_data.get('job_title', '').lower(),
                'location_match': job_seeker_location in job_location if job_seeker_location and job_location else False
            }
            
        except Exception as e:
            print(f"❌ Error when calculating matching score: {e}")
            return {
                'overall_score': 0,
                'matched_skills': [],
                'skill_match': 0,
                'experience_match': False,
                'location_match': False
            }

def get_all_jobs_for_matching():
    """Get all head hunter jobs for matching"""
    try:
        conn = sqlite3.connect('head_hunter_jobs.db')
        c = conn.cursor()
        c.execute("""
            SELECT id, job_title, job_description, main_responsibilities, required_skills,
                   client_company, industry, work_location, work_type, company_size,
                   employment_type, experience_level, visa_support,
                   min_salary, max_salary, currency, benefits
            FROM head_hunter_jobs
            WHERE job_valid_until >= date('now')
        """)
        jobs = c.fetchall()
        conn.close()
        return jobs
    except Exception as e:
        st.error(f"Failed to get job positions: {e}")
        return []

def get_all_job_seekers():
    """Get all job seekers information"""
    try:
        conn = sqlite3.connect('job_seeker.db')
        c = conn.cursor()
        c.execute("""
            SELECT
                id,
                education_level as education,
                work_experience as experience,
                hard_skills as skills,
                industry_preference as target_industry,
                location_preference as target_location,
                salary_expectation as expected_salary,
                university_background as current_title,
                major,
                languages,
                certificates,
                soft_skills,
                project_experience,
                benefits_expectation
            FROM job_seekers
        """)
        seekers = c.fetchall()
        conn.close()

        # Change the structure to match the expected output
        formatted_seekers = []
        for seeker in seekers:
            # Create a virtual name field (using education background + major)
            virtual_name = f"求职者#{seeker[0]} - {seeker[1]}"

            formatted_seekers.append((
                seeker[0],  # id
                virtual_name,  # name (constructed)
                seeker[3] or "",  # skills (hard_skills)
                seeker[2] or "",  # experience (work_experience)
                seeker[1] or "",  # education (education_level)
                seeker[8] or "",  # target_position (major)
                seeker[4] or "",  # target_industry (industry_preference)
                seeker[5] or "",  # target_location (location_preference)
                seeker[6] or "",  # expected_salary (salary_expectation)
                seeker[7] or ""   # current_title (university_background)
            ))

        return formatted_seekers
    except Exception as e:
        st.error(f"Failed to get job seekers: {e}")
        return []
    
def analyze_match_simple(job_data, seeker_data):
    """Simple match analysis between job and seeker"""
    match_score = 50  # Basic Score

    # Skills matching
    job_skills = str(job_data[4]).lower()
    seeker_skills = str(seeker_data[2]).lower()
    skill_match = len(set(job_skills.split()) & set(seeker_skills.split())) / max(len(job_skills.split()), 1)
    match_score += skill_match * 20

    # Experience matching
    experience_map = {"fresh graduate": 0, "1-3 years": 1, "3-5 years": 2, "5-10 years": 3, "10+ years": 4}
    #experience_map = {"应届": 0, "1-3年": 1, "3-5年": 2, "5-10年": 3, "10年以上": 4}
    job_exp = job_data[11]
    seeker_exp = seeker_data[3]

    if job_exp in experience_map and seeker_exp in experience_map:
        exp_diff = abs(experience_map[job_exp] - experience_map[seeker_exp])
        match_score -= exp_diff * 5

    # Industry matching
    job_industry = str(job_data[6]).lower()
    seeker_industry = str(seeker_data[6]).lower()
    if job_industry in seeker_industry or seeker_industry in job_industry:
        match_score += 10

    # Location matching
    job_location = str(job_data[8]).lower()
    seeker_location = str(seeker_data[7]).lower()
    if job_location in seeker_location or seeker_location in job_location:
        match_score += 5

    match_score = max(0, min(100, match_score))

    # Analyze based on score
    if match_score >= 80:
        strengths = ["High skill match", "Experience meets requirements", "Strong industry relevance"]
        #strengths = ["技能高度匹配", "经验符合要求", "行业相关性强"]
        gaps = []
        recommendation = "Highly recommend for interview"
        #recommendation = "强烈推荐面试"
    elif match_score >= 60:
        strengths = ["Core skills match", "Basic experience aligns"]
        #strengths = ["核心技能匹配", "基础经验符合"]
        gaps = ["Some skills need improvement", "Slight experience gap"]
        #gaps = ["部分技能需要提升", "经验略有差距"]
        recommendation = "Recommend further communication"
        #recommendation = "推荐进一步沟通"
    else:
        strengths = ["Has relevant background"]
        #strengths = ["有相关背景"]
        gaps = ["Low skill match", "Experience does not meet requirements"]
        #gaps = ["技能匹配度较低", "经验要求不符"]
        recommendation = "Further evaluation needed"
        #recommendation = "需要进一步评估"

    return {
        "match_score": int(match_score),
        "key_strengths": strengths,
        "potential_gaps": gaps,
        "recommendation": recommendation,
        "salary_match": "Good" if match_score > 70 else "Average",
        #"salary_match": "良好" if match_score > 70 else "一般",
        "culture_fit": "High" if match_score > 75 else "Medium"
        #"culture_fit": "高" if match_score > 75 else "中"
    }

def show_match_statistics():
    """Show match statistics"""
    st.header("📊 Match Statistics")

    jobs = get_all_jobs_for_matching()
    seekers = get_all_job_seekers()

    if not jobs or not seekers:
        st.info("No statistics data available")
        return

    # Industry distribution
    st.subheader("🏭 Industry Distribution")
    industry_counts = {}
    for job in jobs:
        industry = job[6] if job[6] else "Not Specified"
        industry_counts[industry] = industry_counts.get(industry, 0) + 1

    for industry, count in industry_counts.items():
        percentage = (count / len(jobs)) * 100
        st.write(f"• **{industry}:** {count} Positions ({percentage:.1f}%)")

    # Experience Level Distribution
    st.subheader("🎯 Experience Level Distribution")
    experience_counts = {}
    for job in jobs:
        experience = job[11] if job[11] else "Not Specified"
        experience_counts[experience] = experience_counts.get(experience, 0) + 1

    for exp, count in experience_counts.items():
        st.write(f"• **{exp}:** {count} Positions")

def show_instructions():
    """Display usage instructions"""
    st.header("📖 Instructions")

    st.info("""
    **Recruitment Match Instructions:**

    1. **Select Position**: Choose a position from the positions published by the headhunter module
    2. **Set Conditions**: Adjust the minimum match score and display count
    3. **Start Matching**: The system will automatically analyze the match between all job seekers and the position
    4. **View Results**: View detailed match analysis report
    5. **Take Action**: Contact candidates, schedule interviews

    **Matching Algorithm Based on:**
    • Skill Match (Hard Skills)
    • Experience Fit (Work Experience Years)
    • Industry Relevance (Industry Preferences)
    • Location Match (Work Location Preferences)
    • Comprehensive Assessment Analysis

    **Data Sources:**
    • Position Information: Positions published by Head Hunter module
    • Job Seeker Information: Information filled in Job Seeker page
    """)


def get_jobs_for_interview():
    """Get available positions for interviews"""
    try:
        conn = sqlite3.connect('head_hunter_jobs.db')
        c = conn.cursor()
        c.execute("""
            SELECT id, job_title, job_description, main_responsibilities, required_skills,
                   client_company, industry, experience_level
            FROM head_hunter_jobs
            WHERE job_valid_until >= date('now')
        """)
        jobs = c.fetchall()
        conn.close()
        return jobs
    except Exception as e:
        st.error(f"Failed to get positions: {e}")
        return []


def get_job_seeker_profile():
    """Get current job seeker information"""
    try:
        conn = sqlite3.connect('job_seeker.db')
        c = conn.cursor()
        c.execute("""
            SELECT education_level, work_experience, hard_skills, soft_skills,
                   project_experience
            FROM job_seekers
            ORDER BY id DESC
            LIMIT 1
        """)
        profile = c.fetchone()
        conn.close()
        return profile
    except Exception as e:
        st.error(f"Failed to get job seeker information: {e}")
        return None

def initialize_interview_session(job_data):
    """Initialize interview session"""
    if 'interview' not in st.session_state:
        st.session_state.interview = {
            'job_id': job_data[0],
            'job_title': job_data[1],
            'company': job_data[5],
            'current_question': 0,
            'total_questions': 2,
            'questions': [],
            'answers': [],
            'scores': [],
            'completed': False,
            'summary': None
        }

def generate_interview_question(job_data, seeker_profile, previous_qa=None):
    """Generate interview questions using Azure OpenAI"""
    try:
        # Check if API keys are configured
        is_configured, error_msg = Config.check_azure_credentials()
        if not is_configured:
            return f"Error: {error_msg}"
        
        client = AzureOpenAI(
            azure_endpoint=Config.AZURE_ENDPOINT,
            api_key=Config.AZURE_API_KEY,
            api_version=Config.AZURE_API_VERSION
        )

        # Prepare position information
        job_info = f"""
Position Title: {job_data[1]}
Company: {job_data[5]}
Industry: {job_data[6]}
Experience Requirement: {job_data[7]}
Job Description: {job_data[2]}
Main Responsibilities: {job_data[3]}
Required Skills: {job_data[4]}
        """

        # Prepare job seeker information
        seeker_info = ""
        if seeker_profile:
            seeker_info = f"""
Job Seeker Background:
- Education: {seeker_profile[0]}
- Experience: {seeker_profile[1]}
- Hard Skills: {seeker_profile[2]}
- Soft Skills: {seeker_profile[3]}
- Project Experience: {seeker_profile[4]}
            """

        # Build prompt
        if previous_qa:
            prompt = f"""
As a professional interviewer, please continue the interview based on the following information:

【Position Information】
{job_info}

【Job Seeker Information】
{seeker_info}

【Previous Q&A】
Question: {previous_qa['question']}
Answer: {previous_qa['answer']}

Based on the job seeker's previous answer, please ask a relevant follow-up question. The question should:
1. Deeply explore key points from the previous answer
2. Assess the job seeker's thinking depth and professional abilities
3. Be closely related to position requirements

Please only return the question content, without additional explanations.
            """
        else:
            prompt = f"""
As a professional interviewer, please design an interview question for the following position:

【Position Information】
{job_info}

【Job Seeker Information】
{seeker_info}

Please ask a professional interview question that should:
1. Assess core abilities related to the position
2. Examine the job seeker's experience and skills
3. Have appropriate challenge level
4. Can be behavioral, technical, or situational questions

Please only return the question content, without additional explanations.
            """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional recruitment interviewer, skilled at asking targeted interview questions to assess candidates' abilities and suitability."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.8,
            max_tokens=500
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        return f"AI question generation failed: {str(e)}"
    
def evaluate_answer(question, answer, job_data):
    """Evaluate job seeker's answer"""
    try:
        # Check if API keys are configured
        is_configured, error_msg = Config.check_azure_credentials()
        if not is_configured:
            return f'{{"error": "{error_msg}"}}'
        
        client = AzureOpenAI(
            azure_endpoint=Config.AZURE_ENDPOINT,
            api_key=Config.AZURE_API_KEY,
            api_version=Config.AZURE_API_VERSION
        )

        prompt = f"""
Please evaluate the following interview answer:

【Position Information】
Position: {job_data[1]}
Company: {job_data[5]}
Requirements: {job_data[4]}

【Interview Question】
{question}

【Job Seeker Answer】
{answer}

Please evaluate and provide scores (0-10 points) from the following dimensions:
1. Relevance and accuracy of the answer
2. Professional knowledge and skills demonstrated
3. Communication expression and logic
4. Match with position requirements

Please return evaluation results in the following JSON format:
{{
    "score": score,
    "feedback": "Specific feedback and suggestions",
    "strengths": ["Strength1", "Strength2"],
    "improvements": ["Improvement suggestion1", "Improvement suggestion2"]
}}
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional interview evaluation expert, capable of objectively assessing the quality of interview answers."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=800
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        return f'{{"error": "Evaluation failed: {str(e)}"}}'

def generate_final_summary(interview_data, job_data):
    """Generate final interview summary"""
    try:
        # Check if API keys are configured
        is_configured, error_msg = Config.check_azure_credentials()
        if not is_configured:
            return f'{{"error": "{error_msg}"}}'
        
        client = AzureOpenAI(
            azure_endpoint=Config.AZURE_ENDPOINT,
            api_key=Config.AZURE_API_KEY,
            api_version=Config.AZURE_API_VERSION
        )

        # Prepare all Q&A records
        qa_history = ""
        for i, (q, a, score_data) in enumerate(zip(
            interview_data['questions'],
            interview_data['answers'],
            interview_data['scores']
        )):
            qa_history += f"""
Question {i+1}: {q}
Answer: {a}
Score: {score_data.get('score', 'N/A')}
Feedback: {score_data.get('feedback', '')}
            """


        prompt = f"""
Please generate a comprehensive summary report for the following interview:

【Position Information】
Position: {job_data[1]}
Company: {job_data[5]}
Requirements: {job_data[4]}

【Interview Q&A Records】
{qa_history}

Please provide:
1. Overall performance score (0-100 points)
2. Core strengths analysis
3. Areas needing improvement
4. Match assessment for this position
5. Specific improvement suggestions

Please return in the following JSON format:
{{
    "overall_score": overall_score,
    "summary": "Overall evaluation summary",
    "key_strengths": ["Strength1", "Strength2", "Strength3"],
    "improvement_areas": ["Improvement area1", "Improvement area2", "Improvement area3"],
    "job_fit": "High/Medium/Low",
    "recommendations": ["Recommendation1", "Recommendation2", "Recommendation3"]
}}
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional career advisor, capable of providing comprehensive interview performance analysis and career development suggestions."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=1000
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        return f'{{"error": "Summary generation failed: {str(e)}"}}'

def ai_interview_page():
    """AI Interview Page"""
    st.title("🤖 AI Mock Interview")

    # Get position information
    jobs = get_jobs_for_interview()
    seeker_profile = get_job_seeker_profile()

    if not jobs:
        st.warning("❌ No available position information, please first publish positions in the headhunter module")
        return

    if not seeker_profile:
        st.warning("❌ Please first fill in your information on the Job Seeker page")
        return

    st.success("🎯 Select the position you want to interview for to start the mock interview")

    # Select position
    job_options = {f"#{job[0]} {job[1]} - {job[5]}": job for job in jobs}
    selected_job_key = st.selectbox("Select Interview Position", list(job_options.keys()))
    selected_job = job_options[selected_job_key]

    # Display position information
    with st.expander("📋 Position Information", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Position:** {selected_job[1]}")
            st.write(f"**Company:** {selected_job[5]}")
            st.write(f"**Industry:** {selected_job[6]}")
        with col2:
            st.write(f"**Experience Requirement:** {selected_job[7]}")
            st.write(f"**Skill Requirements:** {selected_job[4][:100]}...")

    # Initialize interview session
    initialize_interview_session(selected_job)
    interview = st.session_state.interview

    # Start/continue interview
    if not interview['completed']:
        if interview['current_question'] == 0:
            if st.button("🚀 Start Mock Interview", type="primary", use_container_width=True):
                # Generate first question
                with st.spinner("AI is preparing interview questions..."):
                    first_question = generate_interview_question(selected_job, seeker_profile)
                    if not first_question.startswith("AI question generation failed"):
                        interview['questions'].append(first_question)
                        interview['current_question'] = 1
                        st.rerun()
                    else:
                        st.error(first_question)

        # Display current question
        if interview['current_question'] > 0 and interview['current_question'] <= interview['total_questions']:
            st.subheader(f"❓ Question {interview['current_question']}/{interview['total_questions']}")
            st.info(interview['questions'][-1])

            # Answer input
            answer = st.text_area("Your Answer:", height=150,
                                placeholder="Please describe your answer in detail...",
                                key=f"answer_{interview['current_question']}")


            if st.button("📤 Submit Answer", type="primary", use_container_width=True):
                if answer.strip():
                    with st.spinner("AI is evaluating your answer..."):
                        # Evaluate current answer
                        evaluation = evaluate_answer(
                            interview['questions'][-1],
                            answer,
                            selected_job
                        )

                        try:
                            eval_data = json.loads(evaluation)
                            if 'error' not in eval_data:
                                # Save answer and evaluation
                                interview['answers'].append(answer)
                                interview['scores'].append(eval_data)

                                # Check if all questions are completed
                                if interview['current_question'] == interview['total_questions']:
                                    # Generate final summary
                                    with st.spinner("AI is generating interview summary..."):
                                        summary = generate_final_summary(interview, selected_job)
                                        try:
                                            summary_data = json.loads(summary)
                                            interview['summary'] = summary_data
                                            interview['completed'] = True
                                        except (json.JSONDecodeError, KeyError, TypeError):
                                            interview['summary'] = {"error": "Summary parsing failed"}
                                            interview['completed'] = True
                                else:
                                    # Generate next question
                                    previous_qa = {
                                        'question': interview['questions'][-1],
                                        'answer': answer
                                    }
                                    next_question = generate_interview_question(
                                        selected_job, seeker_profile, previous_qa
                                    )
                                    if not next_question.startswith("AI question generation failed"):
                                        interview['questions'].append(next_question)
                                        interview['current_question'] += 1
                                    else:
                                        st.error(next_question)

                                st.rerun()
                            else:
                                st.error(eval_data['error'])
                        except json.JSONDecodeError:
                            st.error("Evaluation result parsing failed")
                else:
                    st.warning("Please enter your answer")

            # Display progress
            progress = interview['current_question'] / interview['total_questions']
            st.progress(progress)
            st.write(f"Progress: {interview['current_question']}/{interview['total_questions']} questions")

    # Display interview results
    if interview['completed'] and interview['summary']:
        st.subheader("🎯 Interview Summary Report")

        summary = interview['summary']

        if 'error' in summary:
            st.error(summary['error'])
        else:
            # Overall score
            col1, col2, col3 = st.columns(3)
            with col1:
                score = summary.get('overall_score', 0)
                st.metric("Overall Score", f"{score}/100")
            with col2:
                st.metric("Job Fit", summary.get('job_fit', 'N/A'))
            with col3:
                st.metric("Questions Answered", f"{len(interview['answers'])}/{interview['total_questions']}")

            # Overall evaluation
            st.write("### 📊 Overall Evaluation")
            st.info(summary.get('summary', ''))

            # Core strengths
            st.write("### ✅ Core Strengths")
            strengths = summary.get('key_strengths', [])
            for strength in strengths:
                st.write(f"🎯 {strength}")

            # Improvement areas
            st.write("### 📈 Improvement Suggestions")
            improvements = summary.get('improvement_areas', [])
            for improvement in improvements:
                st.write(f"💡 {improvement}")

            # Detailed recommendations
            st.write("### 🎯 Career Development Recommendations")
            recommendations = summary.get('recommendations', [])
            for rec in recommendations:
                st.write(f"🌟 {rec}")

            # Detailed Q&A records
            with st.expander("📝 View Detailed Q&A Records"):
                for i, (question, answer, score_data) in enumerate(zip(
                    interview['questions'],
                    interview['answers'],
                    interview['scores']
                )):
                    st.write(f"#### Question {i+1}")
                    st.write(f"**Question:** {question}")
                    st.write(f"**Answer:** {answer}")
                    if isinstance(score_data, dict):
                        st.write(f"**Score:** {score_data.get('score', 'N/A')}/10")
                        st.write(f"**Feedback:** {score_data.get('feedback', '')}")
                    st.markdown("---")

            # Restart interview
            if st.button("🔄 Restart Interview", use_container_width=True):
                del st.session_state.interview
                st.rerun()