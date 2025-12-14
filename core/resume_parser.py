"""
Resume parsing and profile extraction functionality.

This module provides:
- ResumeParser: Extract text from PDF, DOCX, and TXT files
- GPT4JobRoleDetector: AI-based skill and role detection
- Profile extraction functions
- Streamlit-integrated file extraction

Consolidated from:
- Original core/resume_parser.py
- modules/resume_upload/file_extraction.py
- modules/resume_upload/profile_extraction.py
"""

import re
import json
import requests
from typing import Dict, Optional, Tuple

# Lazy imports for document processing
_PyPDF2 = None
_Document = None


def _get_pypdf2():
    """Lazy load PyPDF2."""
    global _PyPDF2
    if _PyPDF2 is None:
        import PyPDF2
        _PyPDF2 = PyPDF2
    return _PyPDF2


def _get_docx_document():
    """Lazy load python-docx Document."""
    global _Document
    if _Document is None:
        from docx import Document
        _Document = Document
    return _Document


# ============================================================================
# RESUME PARSER
# ============================================================================

class ResumeParser:
    """Parse resume from PDF, DOCX, or TXT files.
    
    This class handles text extraction from various document formats.
    Skill extraction is delegated to AI (GPT-4) rather than using hardcoded lists.
    """
    
    def __init__(self):
        pass
    
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text from PDF file object.
        
        Args:
            pdf_file: File-like object containing PDF data
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If PDF cannot be read
        """
        try:
            PyPDF2 = _get_pypdf2()
            # Reset file position to beginning
            pdf_file.seek(0)
            text = ""
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text from DOCX file object.
        
        Args:
            docx_file: File-like object containing DOCX data
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If DOCX cannot be read
        """
        try:
            from docx import Document
            # Reset file position to beginning
            docx_file.seek(0)
            doc = Document(docx_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def extract_text_from_txt(self, txt_file) -> str:
        """Extract text from TXT file object.
        
        Args:
            txt_file: File-like object containing text data
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If TXT cannot be read
        """
        try:
            txt_file.seek(0)
            text = str(txt_file.read(), "utf-8")
            return text
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def extract_text(self, file_obj, filename: str) -> str:
        """Extract text from uploaded file based on extension.
        
        Args:
            file_obj: File-like object
            filename: Original filename with extension
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
        """
        if filename.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(file_obj)
        elif filename.lower().endswith('.docx'):
            return self.extract_text_from_docx(file_obj)
        elif filename.lower().endswith('.txt'):
            return self.extract_text_from_txt(file_obj)
        else:
            raise ValueError("Unsupported file format. Use PDF, DOCX, or TXT.")
    
    def parse_resume(self, file_obj, filename: str) -> Dict:
        """Parse resume and extract raw text only.
        
        Args:
            file_obj: File-like object containing resume
            filename: Original filename with extension
            
        Returns:
            Dictionary with raw_text and metadata
            
        Raises:
            Exception: If resume cannot be parsed or has insufficient content
        """
        try:
            text = self.extract_text(file_obj, filename)
            
            if not text or len(text.strip()) < 50:
                file_type = filename.split('.')[-1].upper() if '.' in filename else 'file'
                raise ValueError(
                    f"Could not extract sufficient text from resume. "
                    f"This may happen if:\n"
                    f"• The {file_type} is scanned/image-based (try a text-based document)\n"
                    f"• The file is corrupted or password-protected\n"
                    f"• The document is mostly empty\n"
                    f"Please try uploading a different format (PDF or DOCX with selectable text)."
                )
            
            resume_data = {
                'raw_text': text,
                'text_length': len(text),
                'word_count': len(text.split()),
                'filename': filename
            }
            
            return resume_data
            
        except Exception as e:
            raise Exception(f"Error parsing resume: {str(e)}")

    def extract_structured_profile(self, resume_text: str, enable_verification: bool = False, config=None) -> Optional[Dict]:
        """Extract structured profile from resume with optional two-pass verification.
        
        This is a method wrapper around the module-level extract_structured_profile function.
        
        Args:
            resume_text: Full resume text
            enable_verification: Whether to run second verification pass
            config: Optional config object
            
        Returns:
            Extracted profile dictionary or None if extraction fails
        """
        return extract_structured_profile(resume_text, enable_verification, config)


# ============================================================================
# GPT-4 JOB ROLE DETECTOR
# ============================================================================

class GPT4JobRoleDetector:
    """Use GPT-4 to detect job roles and extract skills dynamically.
    
    This class uses Azure OpenAI to analyze resumes and extract:
    - Skills (technical, soft, tools, languages, etc.)
    - Job role recommendations
    - Seniority level
    - Job search keywords
    """
    
    def __init__(self, config=None):
        """Initialize the detector.
        
        Args:
            config: Optional config object. If None, will import from config module.
        """
        self._client = None
        self._config = config
        
        if config is None:
            from config import Config
            self._config = Config
        
        self.model = self._config.AZURE_MODEL
    
    @property
    def client(self):
        """Lazy-load AzureOpenAI client only when needed."""
        if self._client is None:
            from openai import AzureOpenAI
            self._client = AzureOpenAI(
                azure_endpoint=self._config.AZURE_ENDPOINT,
                api_key=self._config.AZURE_API_KEY,
                api_version=self._config.AZURE_API_VERSION
            )
        return self._client
    
    def analyze_resume_for_job_roles(self, resume_data: Dict) -> Dict:
        """Analyze resume with GPT-4 - Extract ALL skills dynamically.
        
        Args:
            resume_data: Dictionary with 'raw_text' key containing resume text
            
        Returns:
            Dictionary with extracted skills, role, and analysis
        """
        resume_text = resume_data.get('raw_text', '')[:3000]
        
        system_prompt = """You are an expert career advisor and resume analyst.

Analyze the resume and extract:
1. ALL skills (technical, soft skills, tools, languages, frameworks, methodologies, domain knowledge)
2. Job role recommendations
3. Seniority level
4. SIMPLE job search keywords (for job board APIs)

IMPORTANT for job search:
- Provide a SIMPLE primary role (e.g., "Program Manager" not complex OR/AND queries)
- Keep search keywords SHORT and COMMON
- Avoid complex boolean logic in search queries

Return JSON with this EXACT structure:
{
    "primary_role": "Simple job title (e.g., Program Manager)",
    "simple_search_terms": ["term1", "term2", "term3"],
    "confidence": 0.95,
    "seniority_level": "Junior/Mid-Level/Senior/Lead/Executive",
    "skills": ["skill1", "skill2", "skill3", ...],
    "core_strengths": ["strength1", "strength2", "strength3"],
    "job_search_keywords": ["keyword1", "keyword2"],
    "optimal_search_query": "Simple search string (just the job title)",
    "location_preference": "Detected or 'United States'",
    "industries": ["industry1", "industry2"],
    "alternative_roles": ["role1", "role2", "role3"]
}"""

        user_prompt = f"""Analyze this resume and extract ALL information:

RESUME:
{resume_text}

IMPORTANT - Extract ALL skills including:
- Programming languages (Python, R, SQL, etc.)
- Tools and software (Tableau, Salesforce, Excel, etc.)
- Methodologies (Agile, Scrum, Kanban, etc.)
- Soft skills (Leadership, Communication, etc.)
- Domain expertise (Banking, Finance, Analytics, etc.)
- Technical skills (Data Analysis, Machine Learning, etc.)
- Languages (English, Cantonese, Mandarin, etc.)

For job search, provide SIMPLE terms that would work on LinkedIn/Indeed (not complex boolean queries).

Be thorough and creative!"""

        try:
            # Check if API keys are configured before attempting API call
            is_configured, error_msg = self._config.check_azure_credentials()
            if not is_configured:
                print(f"❌ Configuration Error: {error_msg}")
                fallback = self._fallback_analysis()
                fallback['_error'] = error_msg
                fallback['_analysis_failed'] = True
                return fallback
            
            print("🤖 Calling GPT-4 for resume analysis...")
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )
            
            ai_analysis = json.loads(response.choices[0].message.content)
            print(f"✅ GPT-4 analysis complete! Found {len(ai_analysis.get('skills', []))} skills")
            
            # Validate that we got meaningful data
            if not ai_analysis.get('primary_role') or ai_analysis.get('primary_role') == 'Professional':
                ai_analysis['_analysis_incomplete'] = True
            
            return ai_analysis
            
        except Exception as e:
            error_msg = str(e)
            print(f"❌ GPT-4 Error: {error_msg}")
            fallback = self._fallback_analysis()
            fallback['_error'] = error_msg
            fallback['_analysis_failed'] = True
            return fallback
    
    def _fallback_analysis(self) -> Dict:
        """Fallback if GPT-4 fails - returns empty strings so user can fill in manually."""
        return {
            "primary_role": "",
            "simple_search_terms": [],
            "confidence": 0.0,
            "seniority_level": "",
            "skills": [],
            "core_strengths": [],
            "job_search_keywords": [],
            "optimal_search_query": "",
            "location_preference": "",
            "industries": [],
            "alternative_roles": [],
            "_analysis_failed": True
        }


# ============================================================================
# PROFILE EXTRACTION FUNCTIONS
# ============================================================================

def extract_relevant_resume_sections(resume_text: str) -> str:
    """Extract Experience and Education sections from resume to reduce token usage.
    
    Args:
        resume_text: Full resume text
        
    Returns:
        Extracted relevant sections (max 2000 chars)
    """
    if not resume_text:
        return ""
    
    experience_keywords = [
        r'experience', r'work experience', r'employment', r'employment history',
        r'professional experience', r'work history', r'career history', r'positions held'
    ]
    education_keywords = [
        r'education', r'academic background', r'academic qualifications',
        r'educational background', r'qualifications', r'degrees'
    ]
    
    lines = resume_text.split('\n')
    relevant_sections = []
    current_section = None
    in_experience = False
    in_education = False
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        line_lower = line_stripped.lower()
        
        if any(re.search(rf'\b{kw}\b', line_lower) for kw in experience_keywords):
            if not in_experience:
                in_experience = True
                in_education = False
                if current_section:
                    relevant_sections.append(current_section)
                current_section = line + '\n'
            continue
        
        if any(re.search(rf'\b{kw}\b', line_lower) for kw in education_keywords):
            if not in_education:
                in_education = True
                if current_section:
                    relevant_sections.append(current_section)
                current_section = line + '\n'
            continue
        
        major_sections = [r'summary', r'objective', r'skills', r'certifications', 
                         r'awards', r'publications', r'projects', r'contact', r'personal']
        if any(re.search(rf'\b{section}\b', line_lower) for section in major_sections):
            if in_experience or in_education:
                if current_section:
                    relevant_sections.append(current_section)
                current_section = None
                in_experience = False
                in_education = False
            continue
        
        if in_experience or in_education:
            if current_section:
                current_section += line + '\n'
    
    if current_section and (in_experience or in_education):
        relevant_sections.append(current_section)
    
    result = '\n'.join(relevant_sections)
    
    # Fallback: look for lines with dates
    if not result or len(result) < 100:
        date_pattern = r'\b(19|20)\d{2}\b|\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}'
        result_lines = []
        for line in lines:
            if re.search(date_pattern, line, re.IGNORECASE):
                result_lines.append(line)
            elif result_lines:
                if len([l for l in result_lines[-3:] if l.strip()]) < 3:
                    result_lines.append(line)
                else:
                    break
        if result_lines:
            result = '\n'.join(result_lines[:50])
    
    if result:
        return result[:2000] if len(result) > 2000 else result
    
    return ""


def extract_structured_profile(resume_text: str, enable_verification: bool = False, config=None) -> Optional[Dict]:
    """Extract structured profile from resume with optional two-pass verification.
    
    Args:
        resume_text: Full resume text
        enable_verification: Whether to run second verification pass
        config: Optional config object
        
    Returns:
        Extracted profile dictionary or None if extraction fails
    """
    if config is None:
        from config import Config
        config = Config
    
    try:
        # Check if API keys are configured
        is_configured, error_msg = config.check_azure_credentials()
        if not is_configured:
            print(f"❌ Configuration Error: {error_msg}")
            return None
        
        from openai import AzureOpenAI
        client = AzureOpenAI(
            azure_endpoint=config.AZURE_ENDPOINT,
            api_key=config.AZURE_API_KEY,
            api_version=config.AZURE_API_VERSION
        )
        
        # FIRST PASS: Initial extraction
        prompt_pass1 = f"""You are an expert at parsing resumes. Extract structured information from the following resume text.

RESUME TEXT:
{resume_text[:6000]}

Please extract and return the following information in JSON format:
{{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL if mentioned",
    "portfolio": "Portfolio/website URL if mentioned",
    "summary": "Professional summary or objective (2-3 sentences)",
    "experience": "Work experience with job titles, companies, dates, and achievements",
    "education": "Education details including degrees, institutions, and graduation dates",
    "skills": "Comma-separated list of technical and soft skills",
    "certifications": "Professional certifications, awards, or achievements"
}}

Important:
- If information is not found, use "N/A" or empty string
- Extract all relevant skills mentioned
- Keep the summary concise but informative
- Return ONLY valid JSON, no additional text"""
        
        print("🤖 Pass 1: Extracting profile information...")
        response_pass1 = client.chat.completions.create(
            model=config.AZURE_MODEL,
            messages=[
                {"role": "system", "content": "You are a resume parser. Extract structured information and return only valid JSON."},
                {"role": "user", "content": prompt_pass1}
            ],
            max_tokens=2000,
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        
        content_pass1 = response_pass1.choices[0].message.content
        profile_data_pass1 = json.loads(content_pass1)
        
        if not enable_verification:
            print("✅ Profile extraction complete (single pass)")
            return profile_data_pass1
        
        # SECOND PASS: Self-correction (optional)
        relevant_sections = extract_relevant_resume_sections(resume_text)
        
        if relevant_sections:
            resume_context = f"""RELEVANT RESUME SECTIONS (Experience and Education only):
{relevant_sections}"""
        else:
            resume_context = f"""RELEVANT RESUME SECTIONS (limited):
{resume_text[:1500]}"""
        
        prompt_pass2 = f"""You are a resume quality checker. Review the extracted profile data against the relevant resume sections and verify accuracy, especially for dates and company names.

{resume_context}

EXTRACTED PROFILE DATA (from first pass):
{json.dumps(profile_data_pass1, indent=2)}

Please review and correct the extracted data, paying special attention to:
1. **Dates** - Verify all employment dates, education dates, and certification dates are accurate
2. **Company Names** - Verify all company/organization names are spelled correctly
3. **Job Titles** - Verify job titles are accurate
4. **Education Institutions** - Verify institution names are correct

Return the corrected profile data in the same JSON format. If everything is correct, return the data as-is.

Return ONLY valid JSON, no additional text."""
        
        print("🔍 Pass 2: Verifying profile data...")
        response_pass2 = client.chat.completions.create(
            model=config.AZURE_MODEL,
            messages=[
                {"role": "system", "content": "You are a resume quality checker. Verify and correct extracted data. Return only valid JSON."},
                {"role": "user", "content": prompt_pass2}
            ],
            max_tokens=2000,
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        content_pass2 = response_pass2.choices[0].message.content
        profile_data_corrected = json.loads(content_pass2)
        print("✅ Profile extraction complete (two-pass verification)")
        return profile_data_corrected
        
    except Exception as e:
        print(f"❌ Profile extraction error: {e}")
        return None


def generate_tailored_resume(user_profile: Dict, job_posting: Dict, 
                             raw_resume_text: str = None, config=None) -> Optional[Dict]:
    """Generate a tailored resume based on user profile and job posting.
    
    Args:
        user_profile: User profile dictionary with name, email, skills, experience, etc.
        job_posting: Job posting dictionary with title, company, description, etc.
        raw_resume_text: Optional original resume text for reference
        config: Optional config object
        
    Returns:
        Generated resume data dictionary or None if generation fails
    """
    if config is None:
        from config import Config
        config = Config
    
    try:
        # Check if API keys are configured
        is_configured, error_msg = config.check_azure_credentials()
        if not is_configured:
            print(f"❌ Configuration Error: {error_msg}")
            return None
        
        from openai import AzureOpenAI
        client = AzureOpenAI(
            azure_endpoint=config.AZURE_ENDPOINT,
            api_key=config.AZURE_API_KEY,
            api_version=config.AZURE_API_VERSION
        )
        
        system_instructions = """You are an expert resume writer with expertise in ATS optimization and career coaching.
Your task is to create a tailored resume by analyzing the job description and adapting the user's profile.
Return ONLY valid JSON - no markdown, no additional text, no code blocks."""

        job_description = f"""JOB POSTING TO MATCH:
Title: {job_posting.get('title', 'N/A')}
Company: {job_posting.get('company', 'N/A')}
Description: {job_posting.get('description', 'N/A')[:3000]}
Required Skills: {', '.join(job_posting.get('skills', [])[:10]) if job_posting.get('skills') else 'N/A'}"""

        structured_profile = f"""STRUCTURED PROFILE:
Name: {user_profile.get('name', 'N/A')}
Email: {user_profile.get('email', 'N/A')}
Phone: {user_profile.get('phone', 'N/A')}
Location: {user_profile.get('location', 'N/A')}
LinkedIn: {user_profile.get('linkedin', 'N/A')}
Summary: {user_profile.get('summary', 'N/A')}
Experience: {user_profile.get('experience', 'N/A')[:2000]}
Education: {user_profile.get('education', 'N/A')}
Skills: {user_profile.get('skills', 'N/A')}
Certifications: {user_profile.get('certifications', 'N/A')}"""

        raw_resume_section = ""
        if raw_resume_text:
            raw_resume_section = f"\n\nORIGINAL RESUME TEXT (for reference):\n{raw_resume_text[:2000]}"

        prompt = f"""{system_instructions}

{job_description}

{structured_profile}{raw_resume_section}

INSTRUCTIONS:
1. Analyze the job posting and identify key skills, technologies, and qualifications needed
2. Tailor the profile to match by:
   - Rewriting the summary to emphasize relevant experience
   - Highlighting skills that match job requirements
   - Rewriting experience bullet points to emphasize relevant achievements
   - Using keywords from the job description for ATS optimization
3. Focus on achievements and measurable results
4. Maintain accuracy - only use information from the provided profile

Return your response as a JSON object with this structure:
{{
  "header": {{
    "name": "Full Name",
    "title": "Professional Title (tailored to job)",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL or empty string"
  }},
  "summary": "2-3 sentence professional summary tailored to the job",
  "skills_highlighted": ["Skill 1", "Skill 2", "Skill 3", ...],
  "experience": [
    {{
      "company": "Company Name",
      "title": "Job Title",
      "dates": "Date Range",
      "bullets": ["Achievement bullet 1...", "Achievement bullet 2..."]
    }}
  ],
  "education": "Education details",
  "certifications": "Certifications and achievements"
}}

Return ONLY the JSON object."""
        
        print("✨ Generating tailored resume...")
        response = client.chat.completions.create(
            model=config.AZURE_MODEL,
            messages=[
                {"role": "system", "content": system_instructions},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000,
            temperature=0.7,
            response_format={"type": "json_object"}
        )
        
        content = response.choices[0].message.content
        resume_data = json.loads(content)
        print("✅ Tailored resume generated!")
        return resume_data
        
    except Exception as e:
        print(f"❌ Resume generation error: {e}")
        return None


# ============================================================================
# STREAMLIT FILE EXTRACTION (from modules/resume_upload/file_extraction.py)
# ============================================================================

def extract_text_from_resume(uploaded_file) -> Optional[str]:
    """Extract text from uploaded resume file (PDF, DOCX, or TXT).
    
    This is the Streamlit-integrated version that works with uploaded file objects.
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        Extracted text or None if extraction fails
    """
    import streamlit as st
    
    try:
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type == 'pdf':
            PyPDF2 = _get_pypdf2()
            uploaded_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            
            # Check if we got any text
            if not text or len(text.strip()) < 50:
                st.warning(
                    "⚠️ Could not extract text from PDF. This may happen if:\n"
                    "• The PDF is scanned/image-based\n"
                    "• The PDF is corrupted or password-protected\n\n"
                    "Please try uploading a DOCX file or a PDF with selectable text."
                )
                return None
            return text
        
        elif file_type == 'docx':
            Document = _get_docx_document()
            uploaded_file.seek(0)
            doc = Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Check if we got any text
            if not text or len(text.strip()) < 50:
                st.warning(
                    "⚠️ Could not extract sufficient text from DOCX. "
                    "The document may be mostly empty or use non-standard formatting."
                )
                return None
            return text
        
        elif file_type == 'txt':
            uploaded_file.seek(0)
            text = str(uploaded_file.read(), "utf-8")
            return text
        
        else:
            st.error(f"Unsupported file type: {file_type}. Please upload PDF, DOCX, or TXT.")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text from resume: {e}")
        return None


# ============================================================================
# STREAMLIT PROFILE EXTRACTION (from modules/resume_upload/profile_extraction.py)
# ============================================================================

def _get_utils_for_profile():
    """Lazy load utility functions to avoid circular imports"""
    from utils import get_text_generator, api_call_with_retry
    from utils.helpers import _websocket_keepalive
    from utils.config import ENABLE_PROFILE_PASS2
    return {
        'get_text_generator': get_text_generator,
        'api_call_with_retry': api_call_with_retry,
        '_websocket_keepalive': _websocket_keepalive,
        'ENABLE_PROFILE_PASS2': ENABLE_PROFILE_PASS2,
    }


def extract_profile_from_resume(resume_text: str) -> Optional[Dict]:
    """Use Azure OpenAI to extract structured profile information from resume text with two-pass self-correction.
    
    This is the Streamlit-integrated version with UI feedback.
    
    Args:
        resume_text: Full resume text
        
    Returns:
        Extracted profile dictionary or None if extraction fails
    """
    import streamlit as st
    
    try:
        utils = _get_utils_for_profile()
        text_gen = utils['get_text_generator']()
        api_call_with_retry = utils['api_call_with_retry']
        _websocket_keepalive = utils['_websocket_keepalive']
        ENABLE_PROFILE_PASS2 = utils['ENABLE_PROFILE_PASS2']
        
        if text_gen is None:
            st.error("⚠️ Azure OpenAI is not configured. Please configure AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in your Streamlit secrets.")
            return None
        
        # FIRST PASS: Initial extraction
        prompt_pass1 = f"""You are an expert at parsing resumes. Extract structured information from the following resume text.

RESUME TEXT:
{resume_text}

Please extract and return the following information in JSON format:
{{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL if mentioned",
    "portfolio": "Portfolio/website URL if mentioned",
    "summary": "Professional summary or objective (2-3 sentences)",
    "experience": "Work experience in chronological order with job titles, companies, dates, and key achievements (formatted as bullet points)",
    "education": "Education details including degrees, institutions, and graduation dates",
    "skills": "Comma-separated list of technical and soft skills",
    "certifications": "Professional certifications, awards, publications, or other achievements"
}}

Important:
- If information is not found, use "N/A" or empty string
- Format experience with clear job titles, companies, dates, and bullet points for achievements
- Extract all relevant skills mentioned
- Keep the summary concise but informative
- Return ONLY valid JSON, no additional text or markdown"""
        
        payload_pass1 = {
            "messages": [
                {"role": "system", "content": "You are a resume parser. Extract structured information and return only valid JSON."},
                {"role": "user", "content": prompt_pass1}
            ],
            "max_tokens": 2000,
            "temperature": 0.3,
            "response_format": {"type": "json_object"}
        }
        
        _websocket_keepalive("Extracting profile information...")
        
        def make_request_pass1():
            return requests.post(
                text_gen.url,
                headers=text_gen.headers,
                json=payload_pass1,
                timeout=45
            )
        
        response_pass1 = api_call_with_retry(make_request_pass1, max_retries=3)
        
        if not response_pass1 or response_pass1.status_code != 200:
            if response_pass1 and response_pass1.status_code == 429:
                st.error("🚫 Rate limit reached for profile extraction after retries. Please wait a few minutes and try again.")
            else:
                error_detail = response_pass1.text[:200] if response_pass1 and response_pass1.text else "No error details"
                endpoint_info = f"Endpoint: {text_gen.url.split('/deployments')[0]}" if text_gen else "Endpoint: Not configured"
                st.error(f"API Error: {response_pass1.status_code if response_pass1 else 'Unknown'} - {error_detail}\n\n{endpoint_info}")
            return None
        
        result_pass1 = response_pass1.json()
        content_pass1 = result_pass1['choices'][0]['message']['content']
        
        if text_gen.token_tracker and 'usage' in result_pass1:
            usage = result_pass1['usage']
            prompt_tokens = usage.get('prompt_tokens', 0)
            completion_tokens = usage.get('completion_tokens', 0)
            text_gen.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
        
        try:
            profile_data_pass1 = json.loads(content_pass1)
        except json.JSONDecodeError:
            json_match = re.search(r'\{.*\}', content_pass1, re.DOTALL)
            if json_match:
                profile_data_pass1 = json.loads(json_match.group())
            else:
                st.error("Could not parse extracted profile data from first pass. Please try again.")
                return None
        
        # SECOND PASS: Self-correction (optional)
        if not ENABLE_PROFILE_PASS2:
            return profile_data_pass1
        
        relevant_resume_sections = extract_relevant_resume_sections(resume_text)
        
        if relevant_resume_sections:
            resume_context = f"""RELEVANT RESUME SECTIONS (Experience and Education only):
{relevant_resume_sections}"""
        else:
            resume_context = f"""RELEVANT RESUME SECTIONS (limited):
{resume_text[:1500]}"""
        
        prompt_pass2 = f"""You are a resume quality checker. Review the extracted profile data against the relevant resume sections and verify accuracy, especially for dates and company names.

{resume_context}

EXTRACTED PROFILE DATA (from first pass):
{json.dumps(profile_data_pass1, indent=2)}

Please review and correct the extracted data, paying special attention to:
1. **Dates** - Verify all employment dates, education dates, and certification dates are accurate
2. **Company Names** - Verify all company/organization names are spelled correctly
3. **Job Titles** - Verify job titles are accurate
4. **Education Institutions** - Verify institution names are correct

Return the corrected profile data in the same JSON format. If everything is correct, return the data as-is. If corrections are needed, return the corrected version.

Return ONLY valid JSON with this structure:
{{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "linkedin": "LinkedIn URL if mentioned",
    "portfolio": "Portfolio/website URL if mentioned",
    "summary": "Professional summary or objective (2-3 sentences)",
    "experience": "Work experience in chronological order with job titles, companies, dates, and key achievements (formatted as bullet points)",
    "education": "Education details including degrees, institutions, and graduation dates",
    "skills": "Comma-separated list of technical and soft skills",
    "certifications": "Professional certifications, awards, publications, or other achievements"
}}

Return ONLY valid JSON, no additional text or markdown."""
        
        payload_pass2 = {
            "messages": [
                {"role": "system", "content": "You are a resume quality checker. Verify and correct extracted data, especially dates and company names. Return only valid JSON."},
                {"role": "user", "content": prompt_pass2}
            ],
            "max_tokens": 2000,
            "temperature": 0.1,
            "response_format": {"type": "json_object"}
        }
        
        _websocket_keepalive("Verifying profile data...")
        
        def make_request_pass2():
            return requests.post(
                text_gen.url,
                headers=text_gen.headers,
                json=payload_pass2,
                timeout=45
            )
        
        response_pass2 = api_call_with_retry(make_request_pass2, max_retries=3)
        
        if response_pass2 and response_pass2.status_code == 200:
            result_pass2 = response_pass2.json()
            content_pass2 = result_pass2['choices'][0]['message']['content']
            
            if text_gen.token_tracker and 'usage' in result_pass2:
                usage = result_pass2['usage']
                prompt_tokens = usage.get('prompt_tokens', 0)
                completion_tokens = usage.get('completion_tokens', 0)
                text_gen.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
            
            try:
                profile_data_corrected = json.loads(content_pass2)
                return profile_data_corrected
            except json.JSONDecodeError:
                st.warning("⚠️ Self-correction pass failed, using initial extraction. Some details may need manual verification.")
                return profile_data_pass1
        else:
            st.warning("⚠️ Self-correction pass failed, using initial extraction. Some details may need manual verification.")
            return profile_data_pass1
            
    except Exception as e:
        st.error(f"Error extracting profile: {e}")
        return None


def verify_profile_data_pass2(profile_data: Dict, resume_text: str) -> Dict:
    """
    Run Pass 2 verification on existing profile data (Lazy Pass 2).
    
    This function is called on-demand before resume generation to verify
    the accuracy of dates, company names, job titles, and education details.
    
    Args:
        profile_data: Dict containing extracted profile data from Pass 1
        resume_text: Original resume text for cross-reference
        
    Returns:
        Dict with verified/corrected profile data, or original data if verification fails
    """
    import streamlit as st
    
    if not profile_data or not resume_text:
        return profile_data
    
    # Check if already verified in this session
    if st.session_state.get('profile_verified', False):
        return profile_data
    
    try:
        utils = _get_utils_for_profile()
        text_gen = utils['get_text_generator']()
        api_call_with_retry = utils['api_call_with_retry']
        _websocket_keepalive = utils['_websocket_keepalive']
        
        if text_gen is None:
            st.warning("⚠️ AI service unavailable. Using unverified profile data.")
            return profile_data
        
        # Extract relevant sections for verification
        relevant_resume_sections = extract_relevant_resume_sections(resume_text)
        
        if relevant_resume_sections:
            resume_context = f"""RELEVANT RESUME SECTIONS (Experience and Education only):
{relevant_resume_sections}"""
        else:
            resume_context = f"""RELEVANT RESUME SECTIONS (limited):
{resume_text[:1500]}"""
        
        prompt_pass2 = f"""You are a resume quality checker. Review the extracted profile data against the relevant resume sections and verify accuracy, especially for dates and company names.

{resume_context}

EXTRACTED PROFILE DATA (to verify):
{json.dumps(profile_data, indent=2)}

Please review and correct the extracted data, paying special attention to:
1. **Dates** - Verify all employment dates, education dates, and certification dates are accurate
2. **Company Names** - Verify all company/organization names are spelled correctly
3. **Job Titles** - Verify job titles are accurate
4. **Education Institutions** - Verify institution names are correct

Return the corrected profile data in the same JSON format. If everything is correct, return the data as-is. If corrections are needed, return the corrected version.

Return ONLY valid JSON with the same structure as the input. No additional text or markdown."""
        
        payload_pass2 = {
            "messages": [
                {"role": "system", "content": "You are a resume quality checker. Verify and correct extracted data, especially dates and company names. Return only valid JSON."},
                {"role": "user", "content": prompt_pass2}
            ],
            "max_tokens": 2000,
            "temperature": 0.1,
            "response_format": {"type": "json_object"}
        }
        
        _websocket_keepalive("Verifying profile data for resume generation...")
        
        def make_request_pass2():
            return requests.post(
                text_gen.url,
                headers=text_gen.headers,
                json=payload_pass2,
                timeout=45
            )
        
        response_pass2 = api_call_with_retry(make_request_pass2, max_retries=2)
        
        if response_pass2 and response_pass2.status_code == 200:
            result_pass2 = response_pass2.json()
            content_pass2 = result_pass2['choices'][0]['message']['content']
            
            if text_gen.token_tracker and 'usage' in result_pass2:
                usage = result_pass2['usage']
                prompt_tokens = usage.get('prompt_tokens', 0)
                completion_tokens = usage.get('completion_tokens', 0)
                text_gen.token_tracker.add_completion_tokens(prompt_tokens, completion_tokens)
            
            try:
                profile_data_verified = json.loads(content_pass2)
                # Mark as verified to avoid re-verification in same session
                st.session_state.profile_verified = True
                st.success("✅ Profile data verified for accuracy")
                return profile_data_verified
            except json.JSONDecodeError:
                st.warning("⚠️ Verification parsing failed. Using original profile data.")
                return profile_data
        else:
            st.warning("⚠️ Verification service unavailable. Using original profile data.")
            return profile_data
            
    except Exception as e:
        st.warning(f"⚠️ Verification error: {e}. Using original profile data.")
        return profile_data
